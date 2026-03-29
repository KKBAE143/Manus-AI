#!/usr/bin/env python3

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

SOURCE_DIR = Path("clean_manuscript_final")
OUTPUT_DIR = Path("final_clean_manuscript_v2")

DROP_PARAGRAPH_PATTERNS = [
    re.compile(r"Proceed to Disorder\s+\d+", re.I),
    re.compile(r"Ready for Disorder\s+\d+", re.I),
    re.compile(r"Ready for next", re.I),
    re.compile(r"^NEXT$", re.I),
    re.compile(r"next module", re.I),
    re.compile(r"continuing seamlessly", re.I),
    re.compile(r"now continuing", re.I),
    re.compile(r"now we begin", re.I),
    re.compile(r"now we move into", re.I),
    re.compile(r"this module includes", re.I),
    re.compile(r"section completed", re.I),
    re.compile(r"disorder \d+ completed", re.I),
    re.compile(r"completed successfully", re.I),
    re.compile(r"status update", re.I),
    re.compile(r"say \"?Proceed", re.I),
    re.compile(r"would you like", re.I),
    re.compile(r"if you want", re.I),
    re.compile(r"once you confirm", re.I),
    re.compile(r"confirm here", re.I),
    re.compile(r"I'll generate", re.I),
    re.compile(r"I'll prepare", re.I),
    re.compile(r"I'll create", re.I),
    re.compile(r"I'll deliver", re.I),
    re.compile(r"I'll clearly inform you", re.I),
    re.compile(r"copy.*paste.*Word", re.I),
    re.compile(r"save as \.docx", re.I),
    re.compile(r"save as \.pdf", re.I),
    re.compile(r"cover page", re.I),
    re.compile(r"title page", re.I),
    re.compile(r"branding discussion", re.I),
    re.compile(r"layout/style discussion", re.I),
    re.compile(r"Word export discussion", re.I),
    re.compile(r"proposed structure", re.I),
    re.compile(r"planned structure", re.I),
    re.compile(r"next section \(preview\)", re.I),
    re.compile(r"this completes your manual", re.I),
    re.compile(r"auto-link them inside Word", re.I),
    re.compile(r"formatted in Word", re.I),
    re.compile(r"page numbers instructions", re.I),
    re.compile(r"PART \w+ is NOT yet fully complete", re.I),
    re.compile(r"\[PASTE START", re.I),
    re.compile(r"\[PASTE END", re.I),
    re.compile(r"Apply your usual styles", re.I),
    re.compile(r"This module will be practitioner gold", re.I),
    re.compile(r"Fantastic - we're now stepping into", re.I),
    re.compile(r"Now we begin Module", re.I),
    re.compile(r"if you approve", re.I),
    re.compile(r"same premium format", re.I),
    re.compile(r"practitioner-ready format", re.I),
    re.compile(r"fully structured, formatted, and sequentially aligned", re.I),
]

DROP_SENTENCE_PATTERNS = [
    re.compile(r"\bNEXT\b", re.I),
    re.compile(r"Proceed to Disorder\s+\d+.*", re.I),
    re.compile(r"Ready for Disorder\s+\d+.*", re.I),
    re.compile(r"Ready for next.*", re.I),
    re.compile(r"next module.*", re.I),
    re.compile(r"continuing.*", re.I),
    re.compile(r"now continuing.*", re.I),
    re.compile(r"now we begin.*", re.I),
    re.compile(r"now we move into.*", re.I),
    re.compile(r"this module includes.*", re.I),
    re.compile(r"section completed.*", re.I),
    re.compile(r"disorder \d+ completed.*", re.I),
    re.compile(r"completed successfully.*", re.I),
    re.compile(r"status update.*", re.I),
    re.compile(r"when ready.*", re.I),
    re.compile(r"say \"?proceed.*", re.I),
    re.compile(r"buddy\b.*", re.I),
    re.compile(r"would you like.*", re.I),
    re.compile(r"if you want.*", re.I),
    re.compile(r"just say.*", re.I),
    re.compile(r"once you confirm.*", re.I),
    re.compile(r"confirm here.*", re.I),
    re.compile(r"awesome.*", re.I),
    re.compile(r"perfect.*", re.I),
    re.compile(r"fantastic.*", re.I),
    re.compile(r"superb.*", re.I),
    re.compile(r"yes buddy.*", re.I),
    re.compile(r"I'll generate.*", re.I),
    re.compile(r"I'll prepare.*", re.I),
    re.compile(r"I'll create.*", re.I),
    re.compile(r"I'll deliver.*", re.I),
    re.compile(r"I'll clearly inform you.*", re.I),
    re.compile(r"return when done.*", re.I),
    re.compile(r"copy.*paste into Word.*", re.I),
    re.compile(r"paste into Word.*", re.I),
    re.compile(r"use \[PAGE BREAK\].*", re.I),
    re.compile(r"apply your styles.*", re.I),
    re.compile(r"save as \.docx.*", re.I),
    re.compile(r"save as \.pdf.*", re.I),
    re.compile(r"cover page discussion.*", re.I),
    re.compile(r"title page discussion.*", re.I),
    re.compile(r"branding discussion.*", re.I),
    re.compile(r"layout/style discussion.*", re.I),
    re.compile(r"word export discussion.*", re.I),
    re.compile(r"proposed structure.*", re.I),
    re.compile(r"planned structure.*", re.I),
    re.compile(r"next section preview.*", re.I),
    re.compile(r"this completes your manual.*", re.I),
    re.compile(r"auto-link them inside Word.*", re.I),
    re.compile(r"formatted in Word.*", re.I),
    re.compile(r"page numbers instructions.*", re.I),
    re.compile(r"Part \w+ is not yet fully complete.*", re.I),
    re.compile(r"\[PASTE START.*", re.I),
    re.compile(r"\[PASTE END.*", re.I),
    re.compile(r"Apply your usual styles.*", re.I),
    re.compile(r"This module will be practitioner gold.*", re.I),
    re.compile(r"Fantastic - we're now stepping into.*", re.I),
    re.compile(r"Now we begin Module.*", re.I),
    re.compile(r"if you approve.*", re.I),
    re.compile(r"same premium format.*", re.I),
    re.compile(r"practitioner-ready format.*", re.I),
    re.compile(r"fully structured, formatted, and sequentially aligned.*", re.I),
]

DROP_LINE_PATTERNS = [
    re.compile(r"^\[SOURCE PAGE \d+\]$", re.I),
    re.compile(r"^NEXT$", re.I),
]


def iter_blocks(doc: DocxDocumentType):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()


def should_drop_paragraph(text: str) -> bool:
    candidate = normalize(text)
    if not candidate:
        return False
    return any(pattern.search(candidate) for pattern in DROP_PARAGRAPH_PATTERNS)


def clean_text(text: str) -> str:
    lines = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if any(pattern.match(line) for pattern in DROP_LINE_PATTERNS):
            continue
        lines.append(line)

    cleaned = "\n".join(lines)
    sentences = re.split(r"(?<=[.!?])\s+", cleaned)
    kept = []
    for sentence in sentences:
        candidate = normalize(sentence)
        if not candidate:
            continue
        if any(pattern.search(candidate) for pattern in DROP_SENTENCE_PATTERNS):
            continue
        kept.append(candidate)
    return "\n".join(kept).strip()


def copy_table(dst: DocxDocumentType, table: Table) -> None:
    cleaned_rows = []
    for row in table.rows:
        values = [clean_text(cell.text) for cell in row.cells]
        if any(value.strip() for value in values):
            cleaned_rows.append(values)

    if not cleaned_rows:
        return

    new_table = dst.add_table(rows=0, cols=len(cleaned_rows[0]))
    for values in cleaned_rows:
        row = new_table.add_row().cells
        for idx, value in enumerate(values):
            row[idx].text = value


def clean_file(input_path: Path, output_path: Path) -> None:
    src = Document(str(input_path))
    dst = Document()

    for block in iter_blocks(src):
        if isinstance(block, Paragraph):
            if should_drop_paragraph(block.text):
                continue
            cleaned = clean_text(block.text)
            if cleaned:
                dst.add_paragraph(cleaned)
        else:
            copy_table(dst, block)

    dst.save(str(output_path))


def main() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    for part in range(1, 16):
        file_name = f"textbook_part_{part:02d}.docx"
        clean_file(SOURCE_DIR / file_name, OUTPUT_DIR / file_name)
    print("Final artifact removal complete.")


if __name__ == "__main__":
    main()
