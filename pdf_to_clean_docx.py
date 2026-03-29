#!/usr/bin/env python3
"""
pdf_to_clean_docx.py

Convert a large PDF into cleaned DOCX manuscript parts with minimal rewriting.

Features
--------
- Page-by-page extraction using PyMuPDF
- Minimal cleaning: removes chat/UI artifacts while preserving content
- Multi-part DOCX output (safer for very large books)
- Page markers + manifest JSON for traceability
- Simple heading detection heuristics
- Optional hard filters for obvious chat/instruction lines

Install
-------
pip install pymupdf python-docx

Usage
-----
python pdf_to_clean_docx.py \
  --input "tab-1.pdf" \
  --output-dir "output_book" \
  --pages-per-docx 250

Optional:
  --start-page 1
  --end-page 2919
  --keep-page-markers
"""

from __future__ import annotations

import argparse
import json
import re
import statistics
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, cast

import fitz  # PyMuPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.styles.style import ParagraphStyle
from docx.shared import Pt


HARD_REMOVE_PATTERNS = [
    r"^\s*\[SOURCE PAGE \d+\]\s*$",
    r"^=====\s*PAGE\s+\d+\s*=====\s*$",
    r"^\s*Sure[!\.]?\s*$",
    r"^\s*Sure[,!]?\s+here.*$",
    r"^\s*Great question[!\.]?.*$",
    r"^\s*Good question[!\.]?.*$",
    r"^\s*Absolutely[!\.]?\s*$",
    r"^\s*Absolutely[,!]?\s+.*$",
    r"^\s*Of course[!\.]?.*$",
    r"^\s*Certainly[!\.]?.*$",
    r"^\s*Here is\s+.*$",
    r"^\s*Here's\s+.*content.*$",
    r"^\s*Here's\s+.*chapter.*$",
    r"^\s*Here's\s+.*section.*$",
    r"^\s*Here's\s+.*module.*$",
    r"^\s*Here's\s+.*table.*$",
    r"^\s*Here's\s+.*the next.*$",
    r"^\s*Let me now.*$",
    r"^\s*Let me (explain|start|begin|continue|present|summarize|write|provide).*$",
    r"^\s*Let's (now|begin|start|dive|explore|look).*$",
    r"^\s*I will now.*$",
    r"^\s*I'll now.*$",
    r"^\s*I'll (begin|start|cover|explain|present|continue|write|move).*$",
    r"^\s*I've (now|already|just|covered|completed|finished).*$",
    r"^\s*As (I|we) (mentioned|discussed|said|noted|explained|covered).*$",
    r"^\s*As promised.*$",
    r"^\s*Now I'll.*$",
    r"^\s*Now let's.*$",
    r"^\s*Now, let's.*$",
    r"^\s*So basically.*$",
    r"^\s*So, (basically|let me|let's|here|we).*$",
    r"^\s*You (can|should|may|might|will) (see|note|notice|find|observe).*$",
    r"^\s*We've covered.*$",
    r"^\s*We (have now|will now|can now|should|need to).*$",
    r"^\s*This is the (next|final|last|first).*section.*$",
    r"^\s*Moving (on|forward|ahead|to the next).*$",
    r"^\s*In this (section|chapter|module|part|unit|topic),?\s*(I|we|you)?.*$",
    r"^\s*Would you like me to.*$",
    r"^\s*Before I generate.*$",
    r"^\s*Just confirm.*$",
    r"^\s*Once you reply.*$",
    r"^\s*Perfect.*$",
    r"^\s*Excellent.*$",
    r"^\s*Outstanding.*$",
    r"^\s*Done.*$",
    r"^\s*Your wish.*$",
    r"^\s*As your wish.*$",
    r"^\s*If you want, I can.*$",
    r"^\s*Here's what I'll do next.*$",
    r"^\s*Here is what I'll do next.*$",
    r"^\s*Shall I.*$",
    r"^\s*Yess+\s*!*\s*$",
    r"^\s*yes+\s*!*\s*$",
    r"^\s*No worries.*$",
    r"^\s*Buddy.*$",
    r"^\s*Dt\. Kaladhar.*$",
    r"^\s*GPT.*$",
    r"^\s*ChatGPT.*$",
    r"^\s*Sources\s*$",
    r"^\s*Visible:\s*\d+%.*$",
    r"^\s*<<?ImageDisplayed>>?\s*$",
    r"^\s*Page \d+ of \d+\s*$",
    r"^\s*Sound good, buddy\??\s*$",
    r"^\s*If yes, I'll proceed.*$",
    r"^\s*Here's the plan for the next output.*$",
    r"^\s*It will include:\s*$",
    r"^\s*Before I start writing, please confirm.*$",
    r"^\s*Would you like this section titled\s*$",
    r"^\s*Both are professional.*$",
    r"^\s*Which one should I go with\??\s*$",
    r"^\s*Now I'll start building it step-by-step with:\s*$",
    r"^\s*Now, everything is ready for final Word formatting\.?\s*$",
    r"^\s*It'll include:\s*$",
    r"^\s*Elegant section dividers\s*$",
    r"^\s*Indexed headers.*$",
    r"^\s*Color-accented boxes.*$",
    r"^\s*The 'Summary & Quick Recall Sheet' at the end\s*$",
    r"^\s*Word\):\s*$",
    r"^\s*Should it be:\s*$",
    r"^\s*or shorter like:\s*$",
    r"^\s*Which one should I lock in.*$",
    r"^\s*Here's what happens next:\s*$",
    r"^\s*Before I start the compilation pass:\s*$",
    r"^\s*So the next step:.*$",
    r"^\s*Final Professional Sign-off\s*$",
    r"^\s*Preface\s*$",
    r"^\s*Liver & Renal Sections\s*$",
    r"^\s*Summary \+ Quick-Recall Sheet\s*$",
    r"^\s*Next Module Preview\s*$",
    r"^\s*Clinical Nutritionist & (Culinary )?Nutrition Expert\s*$",
    r"^\s*NutriRevive Wellness Studio\s*\|.*$",
    r"^\s*Would you like a dedication page.*$",
    r"^\s*\d+\. I'll merge all confirmed content.*$",
    r"^\s*\d+\. I'll generate both \.docx and \.pdf outputs.*$",
    r"^\s*\d+\. You'll receive a version.*$",
    r"^\s*Shall we begin.*$",
    r"^\s*or simply\s*$",
    r"^\s*everyday dietitians\.\s*$",
    r"^\s*complete on its own\??\s*$",
    r"^\s*and practitioner-friendly nutrition modules ever written!?\s*$",
    r"^\s*Next Module Preview\s*-.*$",
    r"^\s*Final Tagline\s*$",
    r"^\s*Now, I'll start compiling the Word version.*$",
    r"^\s*[^A-Za-z0-9]*Dt\. Kaladhar.*$",
    r"^\s*Next Module:\s*\d+.*$",
    r"^\s*Integration\s*$",
    r"^\s*and structure intact\.\s*$",
    r"^\s*digital release\.\s*$",
    r'^\s*Health Through Food Wisdom\."\s*$',
]

SOFT_REMOVE_CONTAINS = [
    "emoji",
    "ui fragments",
    "conversational confirmations",
    "chat interface",
    "output limit",
    "strict preservation extraction",
    "source preservation check",
    "continuity note",
]

EXTRA_CONVERSATIONAL_HINTS = [
    "or simply",
    "which one should i go with",
    "your wish",
    "perfect",
    "outstanding",
    "here's how the final ending block",
    "next module preview",
    "everything is now locked in",
    "so the next step",
    "would you like a dedication page",
    "before i hit finalize",
    "before i start the compilation pass",
    "i'll generate both .docx and .pdf",
    "you'll receive a version",
    "shall we begin",
    "we're now entering",
    "this chapter is where clinical accuracy meets patient practicality",
    "benchmark for practical clinical nutrition handbooks",
    "clinical nutritionist & culinary nutrition expert",
    "nutrirevive wellness studio",
    "publication-ready",
    "ready-to-publish",
    "ready-to-print",
]

EXTRA_DROP_REGEX = [
    re.compile(r'^\s*".*Diet Therapy.*"\??\s*$', re.I),
    re.compile(r'^\s*".*Management.*"\s*$', re.I),
    re.compile(r'^\s*\d+\.\s+Clinical Overview.*$', re.I),
    re.compile(r'^\s*\d+\.\s+Nutritional Assessment Table.*$', re.I),
    re.compile(r'^\s*\d+\.\s+Dietary Goals.*$', re.I),
    re.compile(r'^\s*\d+\.\s+Step-by-Step Modification Logic.*$', re.I),
    re.compile(r'^\s*\d+\.\s+Counseling Dialogue Examples.*$', re.I),
    re.compile(r'^\s*\d+\.\s+Follow-Up.*$', re.I),
    re.compile(r'^\s*complete on its own\??\s*$', re.I),
    re.compile(r'^\s*everyday dietitians\.\s*$', re.I),
    re.compile(r"^\s*it's complete, polished, and publication-ready.*$", re.I),
    re.compile(r'^\s*clinical nutritionist & culinary nutrition expert\s*$', re.I),
    re.compile(r'^\s*nutrirevive wellness studio.*$', re.I),
    re.compile(r'^\s*\d+\.\s+Quick Reference Summary.*$', re.I),
    re.compile(r'^\s*concise clinical intro.*$', re.I),
    re.compile(r'^\s*dietary goals,\s*$', re.I),
    re.compile(r'^\s*focus nutrients,\s*$', re.I),
    re.compile(r'^\s*food grouping.*$', re.I),
    re.compile(r'^\s*Integration\s*$', re.I),
    re.compile(r'^\s*Every section combines clinical MNT tables.*$', re.I),
    re.compile(r'^\s*"Knowledge becomes power only when applied with purpose\."\s*$', re.I),
    re.compile(r'^\s*and structure intact\.\s*$', re.I),
    re.compile(r'^\s*digital release\.\s*$', re.I),
    re.compile(r'^\s*Health Through Food Wisdom\."\s*$', re.I),
    re.compile(r'^\s*Complications\)\s*$', re.I),
    re.compile(r'^\s*glycemic index, carb exchange, meal spacing.*$', re.I),
    re.compile(r'^\s*dietitian and nutritionist - the Diabetes Mellitus.*$', re.I),
    re.compile(r'^\s*Closing Note\s*$', re.I),
    re.compile(r'^\s*\d+\.\s+Vegetarian \+ Non-Vegetarian Sample Menus.*$', re.I),
    re.compile(r'^\s*\d+\.\s+Encourage meal structuring:.*$', re.I),
    re.compile(r'^\s*\d+\.\s+Highlight visible results:.*$', re.I),
    re.compile(r'^\s*\d+\.\s+Reinforce hydration control.*$', re.I),
    re.compile(r'^\s*\d+\.\s+End each session with one actionable takeaway.*$', re.I),
    re.compile(r'^\s*therapy alignment - making nutrition science instantly usable.*$', re.I),
    re.compile(r'^\s*Hypertension \(DASH approach\) - balancing sodium, potassium, and functional heart\s*$', re.I),
    re.compile(r'^\s*Diabetes Mellitus \(DM\) - carb counting, glycemic index mastery, and food synergy for\s*$', re.I),
    re.compile(r'^\s*Thyroid & PCOD Nutrition - endocrine-metabolic crossover and culinary interventions\s*$', re.I),
    re.compile(r'^\s*foods\s*$', re.I),
    re.compile(r'^\s*insulin response\s*$', re.I),
    re.compile(r'^\s*for hormonal balance\s*$', re.I),
    re.compile(r'^\s*"Case\s+\d+\s*-.*"\??\s*$', re.I),
    re.compile(r'^\s*follow diet\.\s*$', re.I),
    re.compile(r'^\s*cucumber slices"\)\.\s*$', re.I),
    re.compile(r'^\s*classroom\.\s*$', re.I),
]

LIKELY_NONCONTENT_CONTAINS = [
    "would you like me to create that file version",
    "once you confirm",
    "i'll now build this into the full file structure",
    "before i generate the word file",
    "just confirm these 3 details",
    "do you prefer the document layout",
    "sound good, buddy",
    "if yes, i'll proceed with the full formatted",
    "here's the plan for the next output",
    "before i start writing, please confirm one small thing",
    "would you like this section titled",
    "which one should i go with",
    "now i'll start building it step-by-step",
    "everything is ready for final word formatting",
    "copy it directly into word",
    "hard copy set",
    "signature manual",
    "same professional-educational, clinic-ready format",
    "this title gives it a clinical touch",
    "the first one emphasizes scientific depth",
    "the second one looks cleaner for clinic use",
    "upcoming cardio-endocrine disorders module",
    "color-accented boxes",
    "indexed headers",
    "most premium and practitioner-friendly nutrition modules ever written",
    "ready-to-print, ready-to-publish version",
    "just confirm one tiny detail for the footer line",
    "the first looks more professional for publication",
    "it's complete, polished, and publication-ready",
    "here's what happens next",
    "before i start the compilation pass",
    "dedication page",
    "i'll merge all confirmed content",
    "i'll generate both .docx and .pdf outputs",
    "you'll receive a version you can open directly in word",
    "final ending block + preview",
    "everything is now locked in",
    "final professional sign-off",
    "professionally formatted word layout",
    "professional, motivating, and reader-engaging",
    "coming up next in the advanced diet modification manual",
    "this upcoming module explores the interlinked pathways",
    "reader message",
    "table of contents",
    "full sections",
    "appendices",
    "acknowledgments",
    "summary & quick recall sheet",
    "endocrine health, connecting:",
    "it's the most essential and applied module",
    "dietitian and nutritionist - the diabetes mellitus",
    "shall we begin",
    "we're now entering the most crucial and career-defining module",
    "this chapter is where clinical accuracy meets patient practicality",
    "you'll learn not only how to plan diets",
    "the skill that sets top clinical dietitians apart",
]

EMOJI_RE = re.compile(
    "[" "\U0001F300-\U0001FAFF" "\u2600-\u27BF" "]+",
    flags=re.UNICODE,
)
WEIRD_BULLETS_RE = re.compile(r"[•●◦▪■◆◇▶►▸▹➤➜➔]+")
MULTISPACE_RE = re.compile(r"[ \t]{2,}")
MULTIBLANK_RE = re.compile(r"\n{3,}")

PERSONAL_DATA_RE = [
    re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}"),
    re.compile(r"(?:ph\.?\s*(?:no\.?|num\.?|number)?\s*[:;]?\s*|phone\s*[:;]\s*|mobile\s*[:;]\s*|contact\s*[:;]\s*)[\+]?\d[\d\s\-]{7,14}", re.I),
    re.compile(r"(?<!\d)\d{10}(?!\d)"),
]

NUMBER_EMOJI_MAP = {
    "1\u20e3": "1.", "2\u20e3": "2.", "3\u20e3": "3.", "4\u20e3": "4.",
    "5\u20e3": "5.", "6\u20e3": "6.", "7\u20e3": "7.", "8\u20e3": "8.",
    "9\u20e3": "9.", "\U0001f51f": "10.",
}
NUMBER_EMOJI_RE = re.compile("|".join(re.escape(k) for k in NUMBER_EMOJI_MAP))

TEXT_FUSION_RE = re.compile(r"([a-z])([A-Z])")
DIGIT_LETTER_RE = re.compile(r"(\d)([A-Za-z])")
LETTER_DIGIT_RE = re.compile(r"([a-z])(\d)")

PIPE_TABLE_RE = re.compile(r"^\s*\|.+\|.+\|\s*$")

CHAT_DROP_PATTERNS = [
    re.compile(r"\bi read the manual\b", re.I),
    re.compile(r"\bfrom my end i want\b", re.I),
    re.compile(r"\bas of now i work\b", re.I),
    re.compile(r"\bsoon gonna leave\b", re.I),
    re.compile(r"\bin that mention what is for what\b", re.I),
    re.compile(r"\bthat'?s exactly the mindset\b", re.I),
    re.compile(r"\byou'?re thinking like a\b", re.I),
    re.compile(r"\blet'?s design this upgrade\b", re.I),
    re.compile(r"\bgreat choice[!.]?\s*starting with\b", re.I),
    re.compile(r"\bnext,\s*would you like me to create\b", re.I),
    re.compile(r"\bwould you like me to create section\b", re.I),
    re.compile(r"\bthis next section will be your secret\b", re.I),
    re.compile(r"\bif you want,?\s*i'?ll now make you\b", re.I),
    re.compile(r"\byess please[!.]?\b", re.I),
    re.compile(r"\bspot on[!.]?\b", re.I),
    re.compile(r"\b-\s*let'?s roll[!.]?\b", re.I),
    re.compile(r"\b-\s*great choice[!.]?\b", re.I),
    re.compile(r"\bkaladhar\s*[-–]\s*we'?re now building\b", re.I),
    re.compile(r"\bit'?ll make you unbeatable\b", re.I),
    re.compile(r"\byour clients will never need a second opinion\b", re.I),
    re.compile(r"\bonce we create this,? you'?ll have\b", re.I),
    re.compile(r"\bword\s*\+\s*pdf version\b", re.I),
    re.compile(r"\bwhich do you prefer\b", re.I),
    re.compile(r"\bconfirm these\s*\d*\s*final details\b", re.I),
    re.compile(r"\bfinal structure[:\s]*['\"]?platinum\b", re.I),
    re.compile(r"\bplanned structure inside the file\b", re.I),
    re.compile(r"\bfront page setup\b.*edition\b", re.I),
    re.compile(r"\blayout selection option\b", re.I),
    re.compile(r"\bformat plan\b", re.I),
    re.compile(r"\bplacement order\b", re.I),
    re.compile(r"\bthis will be your professional toolbook\b", re.I),
    re.compile(r"\bmaster dietitian daily clinical bible\b", re.I),
    re.compile(r"\bdiet modification master manual\b", re.I),
    re.compile(r"\bnavath\s*kaladhar\b", re.I),
    re.compile(r"\bnavathkaladhar\b", re.I),
]

KNOWN_TABLE_HEADERS = [
    re.compile(r"condition\s*/\s*key nutrition focus\s*/\s*supportive nutrients", re.I),
    re.compile(r"foods to include\s*/\s*foods to avoid\s*/\s*clinical tips", re.I),
]

HEADING_PATTERNS = [
    re.compile(r"^\s*SECTION\s+\d+[:\-\s].*$", re.I),
    re.compile(r"^\s*CHAPTER\s+\d+[:\-\s].*$", re.I),
    re.compile(r"^\s*\d+(\.\d+)+\s+.+$"),
    re.compile(r"^\s*[A-Z][A-Z0-9 ,&/\-\(\)]{8,}$"),
]


@dataclass
class PageRecord:
    source_page: int
    cleaned_char_count: int
    kept: bool
    part_number: int | None = None


@dataclass
class SpanRun:
    text: str
    bold: bool
    italic: bool


@dataclass
class LineItem:
    runs: List[SpanRun]
    item_type: str
    font_size: float


@dataclass
class TableItem:
    rows: List[List[str]]
    item_type: str = "table"


ContentItem = Any


LIST_ITEM_RE = re.compile(
    r"^\s*(?:"
    r"[-\u2013\u2014\u2022\u25cf\u25aa\u25ab\u2023\u00b7]\s+"
    r"|[a-zA-Z0-9]{1,3}[.)]\s+"
    r"|\d{1,3}[.)]\s+"
    r")"
)

CASE_HEADER_RE = re.compile(r'^"Case\s+\d+\s*-.*"\??\s*$', re.I)


def _infer_item_type(line_text: str, font_size: float, body_size: float) -> str:
    ratio = font_size / body_size if body_size > 0 else 1.0
    if ratio >= 1.3 or font_size >= 14:
        return "heading-1"
    if ratio >= 1.1 or font_size >= 12:
        return "heading-2"
    if ratio >= 1.05 or font_size >= 11.5:
        return "heading-3"
    if LIST_ITEM_RE.match(line_text):
        return "list-item"
    return "body"


def _line_should_drop(line_text: str) -> bool:
    stripped = line_text.strip()
    if not stripped:
        return True
    if should_drop_line(stripped):
        return True
    if CASE_HEADER_RE.match(stripped):
        return True
    return False


def extract_page_rich(page: Any) -> List[ContentItem]:
    table_rects: List[Tuple[float, float, float, float]] = []
    table_positioned: List[Tuple[float, TableItem]] = []

    try:
        tables = page.find_tables()
        for tbl in tables:
            cells = tbl.extract()
            if not cells:
                continue
            cleaned_rows: List[List[str]] = []
            for row in cells:
                cleaned_row: List[str] = []
                for cell in row:
                    cell_text = normalize_line(str(cell)) if cell is not None else ""
                    if _line_should_drop(cell_text):
                        cell_text = ""
                    cleaned_row.append(cell_text)
                cleaned_rows.append(cleaned_row)
            bbox = tbl.bbox
            rect: Tuple[float, float, float, float] = (float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3]))
            table_rects.append(rect)
            table_positioned.append((rect[1], TableItem(rows=cleaned_rows)))
    except Exception:
        pass

    all_font_sizes: List[float] = []
    try:
        raw_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        for block in raw_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    sz = span.get("size", 11)
                    text = span.get("text", "").strip()
                    if text:
                        all_font_sizes.append(sz)
    except Exception:
        pass

    body_size: float = statistics.median(all_font_sizes) if all_font_sizes else 11.0

    def _block_overlaps_table(block_rect: Tuple[float, float, float, float]) -> bool:
        bx0, by0, bx1, by1 = block_rect
        for t_rect in table_rects:
            tx0, ty0, tx1, ty1 = t_rect
            if bx0 < tx1 and bx1 > tx0 and by0 < ty1 and by1 > ty0:
                return True
        return False

    positioned_items: List[Tuple[float, ContentItem]] = list(table_positioned)

    try:
        raw_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        for block in raw_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            block_rect = block.get("bbox", (0.0, 0.0, 0.0, 0.0))
            if _block_overlaps_table(block_rect):
                continue
            block_y = float(block_rect[1])
            for line in block.get("lines", []):
                line_bbox = line.get("bbox", (0.0, block_y, 0.0, 0.0))
                line_y = float(line_bbox[1])

                runs: List[SpanRun] = []
                dominant_size: float = body_size
                max_sz: float = 0.0

                for span in line.get("spans", []):
                    raw_text = span.get("text", "")
                    cleaned = normalize_line(raw_text)
                    if not cleaned:
                        continue
                    if should_drop_line(cleaned):
                        continue
                    sz = float(span.get("size", body_size))
                    flags = int(span.get("flags", 0))
                    bold = bool(flags & (1 << 4))
                    italic = bool(flags & (1 << 1))
                    runs.append(SpanRun(text=cleaned, bold=bold, italic=italic))
                    if sz > max_sz:
                        max_sz = sz
                        dominant_size = sz

                if not runs:
                    continue

                line_text = "".join(r.text for r in runs).strip()
                if not line_text:
                    continue
                if CASE_HEADER_RE.match(line_text):
                    continue

                item_type = _infer_item_type(line_text, dominant_size, body_size)
                positioned_items.append((line_y, LineItem(
                    runs=runs,
                    item_type=item_type,
                    font_size=dominant_size,
                )))
    except Exception:
        pass

    positioned_items.sort(key=lambda x: x[0])
    return [item for _, item in positioned_items]


def content_items_to_plain_text(items: List[ContentItem]) -> str:
    parts: List[str] = []
    for item in items:
        if isinstance(item, TableItem):
            for row in item.rows:
                row_text = " | ".join(cell for cell in row if cell)
                if row_text:
                    parts.append(row_text)
        elif isinstance(item, LineItem):
            line_text = "".join(r.text for r in item.runs).strip()
            if line_text:
                parts.append(line_text)
    return "\n".join(parts)


def content_items_to_dict(items: List[ContentItem]) -> List[Dict[str, Any]]:
    result: List[Dict[str, Any]] = []
    for item in items:
        if isinstance(item, TableItem):
            result.append({"item_type": "table", "rows": item.rows})
        elif isinstance(item, LineItem):
            result.append({
                "item_type": item.item_type,
                "font_size": item.font_size,
                "runs": [
                    {"text": r.text, "bold": r.bold, "italic": r.italic}
                    for r in item.runs
                ],
            })
    return result


def content_items_from_dict(data: List[Dict[str, Any]]) -> List[ContentItem]:
    result: List[ContentItem] = []
    for d in data:
        if d.get("item_type") == "table":
            result.append(TableItem(rows=d["rows"]))
        else:
            runs = [
                SpanRun(text=r.get("text", ""), bold=bool(r.get("bold", False)), italic=bool(r.get("italic", False)))
                for r in d.get("runs", [])
            ]
            if not runs and d.get("text"):
                runs = [SpanRun(text=d["text"], bold=bool(d.get("bold", False)), italic=bool(d.get("italic", False)))]
            result.append(LineItem(
                runs=runs,
                item_type=d.get("item_type", "body"),
                font_size=float(d.get("font_size", 11.0)),
            ))
    return result


def _replace_number_emoji(m: "re.Match[str]") -> str:
    return NUMBER_EMOJI_MAP.get(m.group(0), m.group(0))


COMMA_ADJACENT_RE = re.compile(r",([A-Za-z])")
TRAILING_PUNCT_RE = re.compile(r"[.!?;:]\s*$")
STARTS_LOWERCASE_OR_DIGIT_RE = re.compile(r"^[a-z0-9\(\[\"]")


def fix_text_fusion(text: str) -> str:
    text = TEXT_FUSION_RE.sub(r"\1 \2", text)
    text = DIGIT_LETTER_RE.sub(r"\1 \2", text)
    text = LETTER_DIGIT_RE.sub(r"\1 \2", text)
    text = COMMA_ADJACENT_RE.sub(r", \1", text)
    return text


def strip_personal_data(line: str) -> str:
    for pat in PERSONAL_DATA_RE:
        line = pat.sub("", line)
    line = MULTISPACE_RE.sub(" ", line).strip()
    return line


def _detect_column_gaps(lines: List[str], min_gap: int = 3) -> List[int]:
    """
    Detect column separator positions in a list of plain-text lines by finding
    positions where ALL lines have a gap of >= min_gap spaces simultaneously.
    Returns a sorted list of column-split positions (0-indexed character offsets).
    """
    if not lines:
        return []
    max_len = max(len(l) for l in lines)
    col_has_char = [False] * max_len
    for line in lines:
        for i, ch in enumerate(line):
            if ch != " ":
                col_has_char[i] = True
    gap_positions: List[int] = []
    in_gap = False
    gap_start = 0
    for i in range(max_len):
        if not col_has_char[i]:
            if not in_gap:
                gap_start = i
                in_gap = True
        else:
            if in_gap:
                if i - gap_start >= min_gap:
                    gap_positions.append(gap_start + (i - gap_start) // 2)
                in_gap = False
    return gap_positions


def reconstruct_column_table(lines: List[str], min_rows: int = 3, min_cols: int = 2, min_gap: int = 4) -> Optional[str]:
    """
    Given a sequence of plain-text lines, attempt to detect column-aligned table
    layout (space-aligned columns as in many PDF text extractions).

    If detected (>= min_rows lines, >= min_cols columns), returns a pipe-delimited
    Markdown-style table string. Otherwise returns None.

    Also deduplicates repeated header rows across page boundaries.
    """
    if len(lines) < min_rows:
        return None
    gaps = _detect_column_gaps(lines, min_gap)
    if len(gaps) < min_cols - 1:
        return None

    split_points = [0] + gaps
    rows: List[List[str]] = []
    for line in lines:
        cells: List[str] = []
        for idx, start in enumerate(split_points):
            end = split_points[idx + 1] if idx + 1 < len(split_points) else len(line)
            cells.append(line[start:end].strip())
        rows.append(cells)

    seen_header: Optional[str] = None
    deduped_rows: List[List[str]] = []
    for row in rows:
        row_key = " | ".join(cell.lower() for cell in row)
        if seen_header is None:
            seen_header = row_key
            deduped_rows.append(row)
        elif row_key == seen_header:
            continue
        else:
            deduped_rows.append(row)

    if len(deduped_rows) < min_rows:
        return None

    table_lines = [" | ".join(row) for row in deduped_rows]
    return "\n".join(table_lines)


def normalize_line(line: str) -> str:
    line = line.replace("\u00a0", " ")
    line = line.replace("\u2018", "'")
    line = line.replace("\u2019", "'")
    line = line.replace("\u201c", '"')
    line = line.replace("\u201d", '"')
    line = line.replace("\u2013", "-")
    line = line.replace("\u2014", "-")
    line = line.replace("\ufe0f", "")
    line = line.replace("\u200d", "")
    line = line.replace("\ufffe", "")
    line = line.replace("\ufffd", "")
    line = NUMBER_EMOJI_RE.sub(_replace_number_emoji, line)
    line = EMOJI_RE.sub("", line)
    line = WEIRD_BULLETS_RE.sub("-", line)
    line = strip_personal_data(line)
    line = MULTISPACE_RE.sub(" ", line).strip()
    return line


def should_drop_line(line: str) -> bool:
    if not line.strip():
        return False

    lower = line.lower().strip()

    for pat in HARD_REMOVE_PATTERNS:
        if re.match(pat, line, flags=re.I):
            return True

    for token in SOFT_REMOVE_CONTAINS:
        if token in lower:
            return True

    for token in EXTRA_CONVERSATIONAL_HINTS:
        if token in lower:
            return True

    for token in LIKELY_NONCONTENT_CONTAINS:
        if token in lower:
            return True

    for pat in EXTRA_DROP_REGEX:
        if pat.match(line):
            return True

    if lower in {"yes", "yess", "perfect", "excellent", "done"}:
        return True

    if PIPE_TABLE_RE.match(line):
        return True

    if re.match(r"^\d{1,2}\.\s*$", line.strip()):
        return True

    for pat in CHAT_DROP_PATTERNS:
        if pat.search(line):
            return True

    return False


def is_heading(line: str) -> bool:
    if len(line) > 140:
        return False
    for pat in HEADING_PATTERNS:
        if pat.match(line):
            return True
    return False


def _is_known_table_header(line: str) -> bool:
    for pat in KNOWN_TABLE_HEADERS:
        if pat.search(line):
            return True
    return False


def _should_join_lines(prev: str, curr: str) -> bool:
    """
    Return True if `curr` is a mid-sentence continuation of `prev` and they should be joined.

    Rules:
    1. prev must not end with sentence-terminal punctuation (.  !  ?  ;  :)
    2. curr must start with lowercase letter, digit, or open bracket — indicating
       it is a continuation of an interrupted sentence from the previous line.
    3. Neither line may be a heading or list item.
    4. The combined length must stay reasonable (<= 400 chars).
    """
    if not prev or not curr:
        return False
    prev_s = prev.rstrip()
    curr_s = curr.strip()
    if not prev_s or not curr_s:
        return False
    if len(prev_s) + 1 + len(curr_s) > 400:
        return False
    if TRAILING_PUNCT_RE.search(prev_s):
        return False
    if not STARTS_LOWERCASE_OR_DIGIT_RE.match(curr_s):
        return False
    if is_heading(prev_s) or is_heading(curr_s):
        return False
    if LIST_ITEM_RE.match(curr_s):
        return False
    return True


def clean_page_text(text: str, seen_table_headers: Optional[set] = None) -> str:
    """
    Clean a single page of extracted PDF text.

    :param text: Raw page text.
    :param seen_table_headers: Optional set shared across page calls to suppress
        repeated table header rows at page boundaries (cross-page dedup).
        Callers that process multi-page sequences should pass the same set object
        for all pages in a chunk.
    """
    raw_lines = text.splitlines()
    kept_lines: List[str] = []
    if seen_table_headers is None:
        seen_table_headers = set()

    candidate_table_blocks: List[Tuple[int, List[str]]] = []
    current_block_start = 0
    current_block_raw: List[str] = []
    for raw_line in raw_lines:
        if raw_line.strip():
            current_block_raw.append(raw_line)
        else:
            if current_block_raw:
                candidate_table_blocks.append((current_block_start, list(current_block_raw)))
            current_block_start += len(current_block_raw) + 1
            current_block_raw = []
    if current_block_raw:
        candidate_table_blocks.append((current_block_start, list(current_block_raw)))

    col_table_replacements: dict = {}
    new_table_header_keys: set = set()
    for block_start, block_raw in candidate_table_blocks:
        if len(block_raw) >= 3:
            table_str = reconstruct_column_table(block_raw, min_rows=3, min_cols=2, min_gap=4)
            if table_str is not None:
                col_table_replacements["\n".join(block_raw)] = table_str
                first_raw = block_raw[0].strip()
                if first_raw:
                    header_key = re.sub(r"\s+", " ", re.sub(r"[|]+", "", first_raw.lower()).strip())
                    new_table_header_keys.add(header_key)

    if col_table_replacements:
        full_text = text
        for orig, replacement in col_table_replacements.items():
            full_text = full_text.replace(orig, replacement)
        raw_lines = full_text.splitlines()

    for raw in raw_lines:
        line = normalize_line(raw)

        if not line:
            kept_lines.append("")
            continue

        if should_drop_line(line):
            continue

        line = fix_text_fusion(line)

        normalized_key = re.sub(r"\s+", " ", re.sub(r"[|]+", "", line.lower()).strip())
        if normalized_key in seen_table_headers:
            continue
        if _is_known_table_header(line):
            seen_table_headers.add(normalized_key)

        if kept_lines and kept_lines[-1] and _should_join_lines(kept_lines[-1], line):
            kept_lines[-1] = kept_lines[-1].rstrip() + " " + line.strip()
            continue

        kept_lines.append(line)

    seen_table_headers.update(new_table_header_keys)

    text = "\n".join(kept_lines)
    text = MULTIBLANK_RE.sub("\n\n", text).strip()

    filtered_blocks: List[str] = []
    seen_headings: set = set()
    for block in text.split("\n\n"):
        block_lines = block.splitlines()
        joined = " ".join(part.strip() for part in block_lines if part.strip())
        if not joined:
            continue

        if re.match(r'^"Case\s+\d+\s*-.*"\??\s*$', joined, flags=re.I):
            continue

        heading_key = re.sub(r"\s+", " ", joined.lower().strip())
        if is_heading(joined) and len(joined) < 120:
            if heading_key in seen_headings:
                continue
            seen_headings.add(heading_key)

        word_count = len(joined.split())
        if word_count < 4 and not is_heading(joined):
            continue

        non_empty_lines = [l for l in block_lines if l.strip()]
        if len(non_empty_lines) >= 3:
            table_str = reconstruct_column_table(non_empty_lines, min_rows=3, min_cols=2, min_gap=4)
            if table_str is not None:
                filtered_blocks.append(table_str)
                continue

        filtered_blocks.append(block)

    text = "\n\n".join(filtered_blocks).strip()
    return text


def _para_hash(text: str) -> str:
    """Normalised fingerprint of a paragraph for near-duplicate detection."""
    return re.sub(r"[^a-z0-9]", "", text.lower())


def deduplicate_paragraphs(paragraphs: List[str], similarity_threshold: float = 0.60) -> List[str]:
    """
    Remove duplicate or near-duplicate paragraphs, keeping only the first occurrence.

    Two paragraphs are considered duplicates when their character-level Jaccard
    similarity (over normalised 4-grams) exceeds similarity_threshold.

    Optimized O(n) for exact duplicates + O(n * W) fuzzy check using a sliding
    window of W recent unique paragraphs (W=120 covers all realistic repeated
    content without the O(n²) full-scan).
    """
    _WINDOW = 120  # recent paragraphs to compare for fuzzy matching

    def _ngrams(s: str, n: int = 4) -> set:
        return {s[i:i+n] for i in range(max(0, len(s) - n + 1))}

    exact_seen: set = set()          # O(1) exact-hash lookup
    window: list = []                # sliding window of (hash, ngram_set)
    result: List[str] = []

    for para in paragraphs:
        if not para.strip():
            result.append(para)
            continue

        h = _para_hash(para)
        if len(h) < 8:
            result.append(para)
            continue

        # Fast exact-duplicate check (O(1))
        if h in exact_seen:
            continue

        # Fuzzy near-duplicate check against recent window (O(W))
        ng = _ngrams(h)
        is_dup = False
        for prev_h, prev_ng in window:
            inter = len(ng & prev_ng)
            union = len(ng | prev_ng)
            if union > 0 and inter / union >= similarity_threshold:
                is_dup = True
                break

        if not is_dup:
            exact_seen.add(h)
            window.append((h, ng))
            if len(window) > _WINDOW:
                window.pop(0)
            result.append(para)

    return result


def extract_pdf_text(pdf_path: Path, start_page: int, end_page: int) -> List[Tuple[int, str]]:
    """
    Extract text from a PDF page range using span-level dict extraction so that
    bold/italic formatting is preserved in the plain-text fallback path.
    """
    doc = fitz.open(pdf_path)
    total_pages = len(doc)

    start_idx = max(0, start_page - 1)
    end_idx = min(total_pages - 1, end_page - 1)

    results: List[Tuple[int, str]] = []
    for i in range(start_idx, end_idx + 1):
        page = doc[i]
        try:
            items = extract_page_rich(page)
            text = content_items_to_plain_text(items)
            if not text.strip():
                text = page.get_text("text")
        except Exception:
            text = page.get_text("text")
        results.append((i + 1, text))
    doc.close()
    return results


def _add_table_borders(tbl: Any) -> None:
    tbl_pr = tbl._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tbl_pr)
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border_el = OxmlElement(f"w:{edge}")
        border_el.set(qn("w:val"), "single")
        border_el.set(qn("w:sz"), "4")
        border_el.set(qn("w:space"), "0")
        border_el.set(qn("w:color"), "000000")
        tbl_borders.append(border_el)
    tbl_pr.append(tbl_borders)


def _shade_table_header_row(tbl: Any, row_idx: int = 0) -> None:
    """Apply light grey shading to the specified table row (default: header row 0)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    row = tbl.rows[row_idx]
    for cell in row.cells:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        tc_pr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True


def _add_body_paragraph_with_spans(doc: Document, text: str, spans: List[dict]) -> None:
    """
    Add a Body Text paragraph, applying bold/italic per-span styling when spans are
    provided. Falls back to plain text if spans do not cover the text.
    """
    try:
        body_style = doc.styles["Body Text"]
    except KeyError:
        body_style = None

    if body_style:
        p = doc.add_paragraph(style="Body Text")
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)

    if not spans:
        p.add_run(text)
        return

    cursor = 0
    for span in spans:
        span_text = span.get("text", "")
        bold = bool(span.get("bold", False))
        italic = bool(span.get("italic", False))
        start = text.find(span_text, cursor)
        if start == -1:
            continue
        if start > cursor:
            p.add_run(text[cursor:start])
        r = p.add_run(span_text)
        r.bold = bold
        r.italic = italic
        cursor = start + len(span_text)
    if cursor < len(text):
        p.add_run(text[cursor:])


def _write_rich_page(doc: Document, items: List[ContentItem]) -> None:
    for item in items:
        if isinstance(item, TableItem):
            if not item.rows:
                continue
            num_cols = max(len(row) for row in item.rows)
            if num_cols == 0:
                continue
            tbl = doc.add_table(rows=len(item.rows), cols=num_cols)
            _add_table_borders(tbl)
            for r_idx, row in enumerate(item.rows):
                for c_idx, cell_text in enumerate(row):
                    if c_idx < num_cols:
                        tbl.cell(r_idx, c_idx).text = cell_text
        elif isinstance(item, LineItem):
            if not item.runs:
                continue
            itype = item.item_type
            if itype == "heading-1":
                p = doc.add_paragraph(style="Heading 1")
            elif itype == "heading-2":
                p = doc.add_paragraph(style="Heading 2")
            elif itype == "heading-3":
                p = doc.add_paragraph(style="Heading 3")
            elif itype == "list-item":
                p = doc.add_paragraph(style="List Bullet")
            else:
                p = doc.add_paragraph()
            for span_run in item.runs:
                r = p.add_run(span_run.text)
                r.bold = span_run.bold
                r.italic = span_run.italic


def build_docx_part(
    part_number: int,
    part_pages: List[Tuple[int, Any]],
    output_path: Path,
    keep_page_markers: bool = False,
) -> None:
    doc = Document()

    style = cast(ParagraphStyle, doc.styles["Normal"])
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    run = title.add_run(f"Cleaned Manuscript Part {part_number:02d}")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("")

    for source_page, page_content in part_pages:
        if keep_page_markers:
            p = doc.add_paragraph()
            r = p.add_run(f"[SOURCE PAGE {source_page}]")
            r.bold = True

        if isinstance(page_content, list):
            _write_rich_page(doc, page_content)
        else:
            text = page_content
            for block in text.split("\n\n"):
                block = block.strip()
                if not block:
                    continue

                lines = block.splitlines()

                if len(lines) == 1 and is_heading(lines[0]):
                    p = doc.add_paragraph(style="Heading 1")
                    p.add_run(lines[0])
                else:
                    if sum(1 for ln in lines if "|" in ln) >= 1:
                        p = doc.add_paragraph()
                        p.add_run("\n".join(lines))
                    else:
                        doc.add_paragraph(" ".join(lines))

        doc.add_page_break()

    doc.save(str(output_path))


ALL_CAPS_HEADING_RE = re.compile(r"^[A-Z][A-Z0-9 ,&/()\-]{5,79}$")


def _is_allcaps_heading(text: str) -> bool:
    stripped = text.strip()
    if not stripped or len(stripped) < 6:
        return False
    if len(stripped) > 80:
        return False
    return bool(ALL_CAPS_HEADING_RE.match(stripped))


def _add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    try:
        p.style = doc.styles["Body Text"]
    except Exception:
        p.style = doc.styles["Normal"]
    p.add_run(text)


def _page_break_before(p) -> None:
    """Attach a page-break-before property to a DOCX paragraph element."""
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pageBreakBefore")
    pb.set(qn("w:val"), "true")
    pPr.append(pb)


def build_docx_from_blocks(
    part_number: int,
    blocks: List[Any],
    output_path: Path,
) -> None:
    """
    Build a DOCX file directly from a list of TransformedBlock-compatible dicts or objects.
    Each block must have: block_type (H1/H2/H3/BODY/LIST/TABLE), text (str).
    TABLE blocks with ' | ' delimiters are converted to DOCX tables.
    ALL-CAPS BODY lines are promoted to Heading 2.
    H1 blocks trigger a page-break-before (not on the very first H1).
    """
    doc = Document()

    style = cast(ParagraphStyle, doc.styles["Normal"])
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    run = title.add_run(f"Transformed Manuscript Part {part_number:02d}")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph("")

    first_h1_seen = False

    for block in blocks:
        btype = block.get("block_type", "BODY") if isinstance(block, dict) else getattr(block, "block_type", "BODY")
        text = block.get("text", "") if isinstance(block, dict) else getattr(block, "text", "")
        spans: List[dict] = block.get("spans", []) if isinstance(block, dict) else getattr(block, "spans", [])
        if not text or not text.strip():
            continue

        text_stripped = MULTIBLANK_RE.sub("\n\n", text.strip()).strip()
        if not text_stripped:
            continue

        if btype == "H1":
            p = doc.add_paragraph(style="Heading 1")
            p.add_run(text_stripped)
            if first_h1_seen:
                _page_break_before(p)
            first_h1_seen = True
        elif btype == "H2":
            p = doc.add_paragraph(style="Heading 2")
            p.add_run(text_stripped)
        elif btype == "H3":
            p = doc.add_paragraph(style="Heading 3")
            p.add_run(text_stripped)
        elif btype == "LIST":
            p = doc.add_paragraph(style="List Bullet")
            clean = re.sub(r"^[-*\u2022]\s*", "", text_stripped)
            p.add_run(clean)
        elif btype == "TABLE":
            rows_raw = [line for line in text.splitlines() if line.strip() and "|" in line]
            if rows_raw:
                sep = " | " if " | " in rows_raw[0] else "|"
                table_rows = [[cell.strip() for cell in row.split(sep)] for row in rows_raw]
                table_rows = [[c for c in row if c] for row in table_rows]
                table_rows = [row for row in table_rows if row]
                if table_rows:
                    num_cols = max(len(r) for r in table_rows)
                    tbl = doc.add_table(rows=len(table_rows), cols=num_cols)
                    _add_table_borders(tbl)
                    for r_idx, row in enumerate(table_rows):
                        for c_idx, cell_text in enumerate(row):
                            if c_idx < num_cols:
                                tbl.cell(r_idx, c_idx).text = cell_text
                        if r_idx == 0:
                            _shade_table_header_row(tbl, row_idx=0)
                else:
                    _add_body_paragraph(doc, text_stripped)
            else:
                _add_body_paragraph(doc, text_stripped)
        else:
            if _is_allcaps_heading(text_stripped):
                p = doc.add_paragraph(style="Heading 2")
                p.add_run(text_stripped)
            elif spans:
                _add_body_paragraph_with_spans(doc, text_stripped, spans)
            else:
                _add_body_paragraph(doc, text_stripped)

    doc.save(str(output_path))


def chunk_pages(pages: List[Tuple[int, str]], pages_per_docx: int) -> List[List[Tuple[int, str]]]:
    parts: List[List[Tuple[int, str]]] = []
    for i in range(0, len(pages), pages_per_docx):
        parts.append(pages[i : i + pages_per_docx])
    return parts


def write_manifest(
    records: List[PageRecord],
    output_dir: Path,
    input_pdf: str,
    start_page: int,
    end_page: int,
    pages_per_docx: int,
) -> None:
    manifest = {
        "input_pdf": input_pdf,
        "start_page": start_page,
        "end_page": end_page,
        "pages_per_docx": pages_per_docx,
        "page_records": [asdict(r) for r in records],
    }
    with (output_dir / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, indent=2, ensure_ascii=False)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, help="Path to input PDF")
    parser.add_argument("--output-dir", required=True, help="Directory for DOCX outputs")
    parser.add_argument(
        "--pages-per-docx",
        type=int,
        default=250,
        help="Source pages per DOCX part",
    )
    parser.add_argument("--start-page", type=int, default=1, help="1-based start page")
    parser.add_argument("--end-page", type=int, default=10_000, help="1-based end page")
    parser.add_argument(
        "--keep-page-markers",
        action="store_true",
        help="Insert [SOURCE PAGE X] markers",
    )
    args = parser.parse_args()

    input_pdf = Path(args.input)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    extracted = extract_pdf_text(input_pdf, args.start_page, args.end_page)

    cleaned_pages: List[Tuple[int, str]] = []
    records: List[PageRecord] = []

    for source_page, raw_text in extracted:
        cleaned = clean_page_text(raw_text)
        keep = bool(cleaned.strip())

        records.append(
            PageRecord(
                source_page=source_page,
                cleaned_char_count=len(cleaned),
                kept=keep,
            )
        )

        if keep:
            cleaned_pages.append((source_page, cleaned))

    parts = chunk_pages(cleaned_pages, args.pages_per_docx)

    for part_number, part_pages in enumerate(parts, start=1):
        for page_num, _ in part_pages:
            for rec in records:
                if rec.source_page == page_num:
                    rec.part_number = part_number
                    break

        output_path = output_dir / f"textbook_part_{part_number:02d}.docx"
        build_docx_part(
            part_number=part_number,
            part_pages=part_pages,
            output_path=output_path,
            keep_page_markers=args.keep_page_markers,
        )

    write_manifest(
        records=records,
        output_dir=output_dir,
        input_pdf=str(input_pdf),
        start_page=args.start_page,
        end_page=min(args.end_page, extracted[-1][0] if extracted else args.start_page),
        pages_per_docx=args.pages_per_docx,
    )

    print(f"Done. Wrote {len(parts)} DOCX part(s) to: {output_dir}")


if __name__ == "__main__":
    main()
