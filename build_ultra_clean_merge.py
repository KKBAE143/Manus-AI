#!/usr/bin/env python3

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

SOURCE_DIR = Path("final_book_merge")
OUTPUT_DIR = Path("ultra_clean_merge")

DROP_PARAGRAPH_PATTERNS = [
    re.compile(r"^Preamble\s*&\s*Purpose of the Manual$", re.I),
    re.compile(r"^Who This Manual Is For$", re.I),
    re.compile(r"^Purpose of this section$", re.I),
    re.compile(r"^Why this section exists$", re.I),
    re.compile(r"^This module teaches.*$", re.I),
    re.compile(r"^This manual helps you.*$", re.I),
    re.compile(r"^place this section before.*$", re.I),
    re.compile(r"^insert this section.*$", re.I),
    re.compile(r"^layout suggestion.*$", re.I),
    re.compile(r"^word formatting.*$", re.I),
    re.compile(r"^color blocks.*$", re.I),
    re.compile(r"^heading style instructions.*$", re.I),
]

DROP_SENTENCE_PATTERNS = [
    re.compile(r"\bPreamble\s*&\s*Purpose of the Manual\b", re.I),
    re.compile(r"\bWho This Manual Is For\b", re.I),
    re.compile(r"\bPurpose of this section\b", re.I),
    re.compile(r"\bWhy this section exists\b", re.I),
    re.compile(r"\bThis module teaches\b.*", re.I),
    re.compile(r"\bThis manual helps you\b.*", re.I),
    re.compile(r"\bpractitioner-ready\b", re.I),
    re.compile(r"\bpolished format\b", re.I),
    re.compile(r"\bpremium edition\b", re.I),
    re.compile(r"\bplatinum edition\b", re.I),
    re.compile(r"\badvanced manual\b", re.I),
    re.compile(r"\bclinic weapon\b", re.I),
    re.compile(r"\bSUPER IMPORTANT\b", re.I),
    re.compile(r"\bultra-fast\b", re.I),
    re.compile(r"\byou will never forget this\b", re.I),
    re.compile(r"\byou'?ll master instantly\b", re.I),
    re.compile(r"\blike a pro\b", re.I),
    re.compile(r"\bplace this section before\b.*", re.I),
    re.compile(r"\binsert this section\b.*", re.I),
    re.compile(r"\blayout suggestion\b.*", re.I),
    re.compile(r"\bword formatting\b.*", re.I),
    re.compile(r"\bcolor blocks\b.*", re.I),
    re.compile(r"\bheading style instructions\b.*", re.I),
]

DROP_LINE_PATTERNS = [
    re.compile(r"^\[PAGE BREAK\]$", re.I),
    re.compile(r"^\[PASTE START\]$", re.I),
    re.compile(r"^\[PASTE END\]$", re.I),
    re.compile(r"^[⭐*\s]+$"),
]

STAR_RE = re.compile(r"[⭐]+")
EMOJI_RE = re.compile(r"[\U0001F300-\U0001FAFF\u2600-\u27BF]+", re.UNICODE)
MULTIBLANK_RE = re.compile(r"\n{3,}")


def iter_blocks(doc: DocxDocumentType):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def normalize(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = STAR_RE.sub("", text)
    text = EMOJI_RE.sub("", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def should_drop_paragraph(text: str) -> bool:
    candidate = normalize(text)
    if not candidate:
        return False
    return any(pattern.search(candidate) for pattern in DROP_PARAGRAPH_PATTERNS)


def clean_text(text: str) -> str:
    text = STAR_RE.sub("", text)
    text = EMOJI_RE.sub("", text)
    lines = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            lines.append("")
            continue
        if any(pattern.match(line) for pattern in DROP_LINE_PATTERNS):
            continue
        lines.append(line)

    text = "\n".join(lines)
    text = MULTIBLANK_RE.sub("\n\n", text).strip()
    pieces = re.split(r"(?<=[.!?])\s+", text)
    kept = []
    for piece in pieces:
        candidate = normalize(piece)
        if not candidate:
            continue
        if any(pattern.search(candidate) for pattern in DROP_SENTENCE_PATTERNS):
            continue
        kept.append(candidate)
    return "\n".join(kept).strip()


def copy_table(dst: DocxDocumentType, table: Table) -> None:
    cleaned_rows = []
    for row in table.rows:
        values = [clean_text(cell.text) for cell in row.cells]
        if any(value.strip() for value in values):
            cleaned_rows.append(values)
    if not cleaned_rows:
        return
    new_table = dst.add_table(rows=0, cols=len(cleaned_rows[0]))
    for values in cleaned_rows:
        row = new_table.add_row().cells
        for idx, value in enumerate(values):
            row[idx].text = value


def clean_file(input_path: Path, output_path: Path) -> None:
    src = Document(str(input_path))
    dst = Document()
    for block in iter_blocks(src):
        if isinstance(block, Paragraph):
            if should_drop_paragraph(block.text):
                continue
            cleaned = clean_text(block.text)
            if cleaned:
                dst.add_paragraph(cleaned)
        else:
            copy_table(dst, block)
    dst.save(str(output_path))


def main() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    for part in range(1, 16):
        name = f"textbook_part_{part:02d}.docx"
        clean_file(SOURCE_DIR / name, OUTPUT_DIR / name)
    print("Ultra clean merge folder ready.")


if __name__ == "__main__":
    main()
