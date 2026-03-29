#!/usr/bin/env python3

import shutil
from pathlib import Path

import final_clean_part01 as cleaner
from docx import Document

OUTPUT_DIR = Path("clean_manuscript_final")
SOURCE_V2 = Path("clean_manuscript_v2")
SOURCE_V3 = Path("clean_manuscript_v3")
ULTRA_CLEAN_PART01 = Path("textbook_part_01_ultra_clean.docx")

PREFERRED_V3 = {"textbook_part_12.docx", "textbook_part_15.docx"}


def clean_docx(input_path: Path, output_path: Path) -> None:
    src = Document(str(input_path))
    dst = Document()

    for block in cleaner.iter_blocks(src):
        if block.__class__.__name__ == "Paragraph":
            if cleaner.should_drop_paragraph(block.text):
                continue
            cleaned = cleaner.clean_text(block.text)
            if cleaned:
                dst.add_paragraph(cleaned)
        else:
            cleaner.copy_table(dst, block)

    dst.save(str(output_path))


def resolve_source(file_name: str) -> Path:
    if file_name in PREFERRED_V3 and (SOURCE_V3 / file_name).exists():
        return SOURCE_V3 / file_name
    return SOURCE_V2 / file_name


def main() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)

    shutil.copyfile(ULTRA_CLEAN_PART01, OUTPUT_DIR / "textbook_part_01.docx")

    for part in range(2, 16):
        file_name = f"textbook_part_{part:02d}.docx"
        src = resolve_source(file_name)
        clean_docx(src, OUTPUT_DIR / file_name)

    print("Final manuscript cleaning complete.")


if __name__ == "__main__":
    main()
