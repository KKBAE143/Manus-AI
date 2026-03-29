#!/usr/bin/env python3

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

SOURCE_DIR = Path("ultra_final_manuscript")
OUTPUT_DIR = Path("final_book_merge")

DROP_PARAGRAPH_PATTERNS = [
    re.compile(r"^WHERE TO PLACE IT:?$", re.I),
    re.compile(r"^Place this section immediately before.*$", re.I),
    re.compile(r"^Pro Tip: In Word.*$", re.I),
    re.compile(r"^add it right before.*$", re.I),
    re.compile(r"^keep it printed.*$", re.I),
    re.compile(r"^laminate it.*$", re.I),
    re.compile(r"^ready for Word paste.*$", re.I),
    re.compile(r"^color block design.*$", re.I),
    re.compile(r"^soft blue headers.*$", re.I),
    re.compile(r"^neutral gray text blocks.*$", re.I),
    re.compile(r"^assign heading styles.*$", re.I),
    re.compile(r"^layout\s*[->]+\s*columns.*$", re.I),
    re.compile(r"^Word formatting instructions.*$", re.I),
]

DROP_SENTENCE_PATTERNS = [
    re.compile(r"\bpractitioner-ready\b", re.I),
    re.compile(r"\bpremium format\b", re.I),
    re.compile(r"\bpolished format\b", re.I),
    re.compile(r"\bhighly practical\b", re.I),
    re.compile(r"\bclinic autopilot\b", re.I),
    re.compile(r"\bevery dietitian must master\b", re.I),
    re.compile(r"\bWHERE TO PLACE IT:?\b", re.I),
    re.compile(r"\bPlace this section immediately before\b.*", re.I),
    re.compile(r"\bPro Tip: In Word\b.*", re.I),
    re.compile(r"\badd it right before\b.*", re.I),
    re.compile(r"\bkeep it printed\b.*", re.I),
    re.compile(r"\blaminate it\b.*", re.I),
    re.compile(r"\bready for Word paste\b.*", re.I),
    re.compile(r"\bcolor block design\b.*", re.I),
    re.compile(r"\bsoft blue headers\b.*", re.I),
    re.compile(r"\bneutral gray text blocks\b.*", re.I),
    re.compile(r"\bassign heading styles\b.*", re.I),
    re.compile(r"\blayout\s*[->]+\s*columns\b.*", re.I),
    re.compile(r"\bWord formatting instructions\b.*", re.I),
]

STAR_ONLY_RE = re.compile(r"^[\s\*⭐]+$")
MULTIBLANK_RE = re.compile(r"\n{3,}")


def iter_blocks(doc: DocxDocumentType):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def normalize(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = text.replace("⭐", "")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def should_drop_paragraph(text: str) -> bool:
    candidate = normalize(text)
    if not candidate or STAR_ONLY_RE.match(text.strip()):
        return True
    return any(pattern.search(candidate) for pattern in DROP_PARAGRAPH_PATTERNS)


def clean_text(text: str) -> str:
    text = text.replace("format Disorder 48 - Recurrent Urinary Tract Infection", "DISORDER 48 - Recurrent Urinary Tract Infection")
    text = text.replace("⭐", "")
    lines = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            lines.append("")
            continue
        if line.lower().startswith("[source page "):
            continue
        if line in {"[PAGE BREAK]", "[PASTE START]", "[PASTE END]"}:
            continue
        if STAR_ONLY_RE.match(line):
            continue
        lines.append(line)

    text = "\n".join(lines)
    text = MULTIBLANK_RE.sub("\n\n", text).strip()
    pieces = re.split(r"(?<=[.!?])\s+", text)
    kept = []
    for piece in pieces:
        candidate = normalize(piece)
        if not candidate:
            continue
        if any(pattern.search(candidate) for pattern in DROP_SENTENCE_PATTERNS):
            continue
        kept.append(candidate)
    return "\n".join(kept).strip()


def copy_table(dst: DocxDocumentType, table: Table) -> None:
    cleaned_rows = []
    for row in table.rows:
        values = [clean_text(cell.text) for cell in row.cells]
        if any(value.strip() for value in values):
            cleaned_rows.append(values)
    if not cleaned_rows:
        return
    new_table = dst.add_table(rows=0, cols=len(cleaned_rows[0]))
    for values in cleaned_rows:
        row = new_table.add_row().cells
        for idx, value in enumerate(values):
            row[idx].text = value


def clean_file(input_path: Path, output_path: Path) -> None:
    src = Document(str(input_path))
    dst = Document()
    for block in iter_blocks(src):
        if isinstance(block, Paragraph):
            if should_drop_paragraph(block.text):
                continue
            cleaned = clean_text(block.text)
            if cleaned:
                dst.add_paragraph(cleaned)
        else:
            copy_table(dst, block)
    dst.save(str(output_path))


def main() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    for part in range(1, 16):
        name = f"textbook_part_{part:02d}.docx"
        clean_file(SOURCE_DIR / name, OUTPUT_DIR / name)
    print("Final book merge folder ready.")


if __name__ == "__main__":
    main()
