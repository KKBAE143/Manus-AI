#!/usr/bin/env python3

import os
import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

INPUT_FOLDER = Path("output_book_markers_200_cleaned")
OUTPUT_FOLDER = Path("clean_manuscript")

OUTPUT_FOLDER.mkdir(exist_ok=True)

ARTIFACT_PATTERNS = [
    r"If you want.*",
    r"Would you like.*",
    r"I'll now.*",
    r"I will now.*",
    r"Let's build.*",
    r"Once you reply.*",
    r"Next module.*",
    r"Superb.*",
    r"I'll generate.*",
    r"I can create.*",
    r"I can prepare.*",
    r"Confirm these.*",
    r"Just confirm.*",
    r"Do you want.*",
    r"Should I include.*",
]


def is_artifact(line: str) -> bool:
    text = line.strip()
    if not text:
        return False

    for pattern in ARTIFACT_PATTERNS:
        if re.match(pattern, text, re.IGNORECASE):
            return True
    return False


def iter_blocks(doc: DocxDocument):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def add_clean_paragraph(dst: Document, text: str) -> None:
    lines = [line.strip() for line in text.splitlines()]
    kept = [line for line in lines if line and not is_artifact(line)]
    if kept:
        dst.add_paragraph("\n".join(kept))


def add_clean_table(dst: Document, table: Table) -> None:
    cleaned_rows = []
    for row in table.rows:
        cleaned_cells = []
        for cell in row.cells:
            lines = [line.strip() for line in cell.text.splitlines()]
            kept = [line for line in lines if line and not is_artifact(line)]
            cleaned_cells.append("\n".join(kept).strip())
        if any(cleaned_cells):
            cleaned_rows.append(cleaned_cells)

    if not cleaned_rows:
        return

    new_table = dst.add_table(rows=0, cols=len(cleaned_rows[0]))
    for row_values in cleaned_rows:
        row = new_table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].text = value


def clean_file(input_path: Path, output_path: Path) -> None:
    src = Document(str(input_path))
    dst = Document()

    for block in iter_blocks(src):
        if isinstance(block, Paragraph):
            add_clean_paragraph(dst, block.text)
        elif isinstance(block, Table):
            add_clean_table(dst, block)

    dst.save(str(output_path))


def main() -> None:
    for file_name in sorted(os.listdir(INPUT_FOLDER)):
        if file_name.endswith(".docx"):
            clean_file(INPUT_FOLDER / file_name, OUTPUT_FOLDER / file_name)

    print("Cleaning complete.")


if __name__ == "__main__":
    main()
