#!/usr/bin/env python3

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

SOURCE_DIR = Path("ultra_clean_merge")
OUTPUT_DIR = Path("final_book_ready")

DROP_PARAGRAPH_PATTERNS = [
    re.compile(r"^SECTION A\s*-.*$", re.I),
    re.compile(r"^SECTION B\s*-.*$", re.I),
    re.compile(r"^FOUNDATION MODULES.*$", re.I),
    re.compile(r"^THERAPEUTIC MODULES.*$", re.I),
    re.compile(r"^PRACTICE MODULES.*$", re.I),
    re.compile(r"^PROFESSIONAL TOOLKIT.*$", re.I),
    re.compile(r"^MODULE LISTS.*$", re.I),
    re.compile(r"^INDEX CREATION TEXT.*$", re.I),
]

DROP_SENTENCE_PATTERNS = [
    re.compile(r"\bshould we add\b.*", re.I),
    re.compile(r"\bthis section fits before\b.*", re.I),
    re.compile(r"\bplace before glossary\b.*", re.I),
    re.compile(r"\bdo you want\b.*", re.I),
    re.compile(r"\blet'?s add\b.*", re.I),
    re.compile(r"\bfits right before\b.*", re.I),
    re.compile(r"\byou'?ll never forget\b.*", re.I),
    re.compile(r"\bsuper important\b.*", re.I),
    re.compile(r"\bclinic weapon\b.*", re.I),
    re.compile(r"\badvanced manual\b.*", re.I),
    re.compile(r"\bpremium edition\b.*", re.I),
    re.compile(r"\binstant mastery\b.*", re.I),
    re.compile(r"\bultimate system\b.*", re.I),
]

STAR_RE = re.compile(r"[⭐]+")
STRAY_SYMBOL_RE = re.compile(r"^[\s*#>~`_^\-]+$")
MULTIBLANK_RE = re.compile(r"\n{3,}")


def iter_blocks(doc: DocxDocumentType):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def normalize(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = STAR_RE.sub("", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def should_drop_paragraph(text: str) -> bool:
    candidate = normalize(text)
    if not candidate or STRAY_SYMBOL_RE.match(text.strip()):
        return True
    return any(pattern.search(candidate) for pattern in DROP_PARAGRAPH_PATTERNS)


def clean_text(text: str) -> str:
    text = STAR_RE.sub("", text)
    lines = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            lines.append("")
            continue
        if STRAY_SYMBOL_RE.match(line):
            continue
        lines.append(line)

    text = "\n".join(lines)
    text = MULTIBLANK_RE.sub("\n\n", text).strip()
    pieces = re.split(r"(?<=[.!?])\s+", text)
    kept = []
    for piece in pieces:
        candidate = normalize(piece)
        if not candidate:
            continue
        if any(pattern.search(candidate) for pattern in DROP_SENTENCE_PATTERNS):
            continue
        kept.append(candidate)
    return "\n".join(kept).strip()


def copy_table(dst: DocxDocumentType, table: Table) -> None:
    cleaned_rows = []
    for row in table.rows:
        values = [clean_text(cell.text) for cell in row.cells]
        if any(value.strip() for value in values):
            cleaned_rows.append(values)
    if not cleaned_rows:
        return
    new_table = dst.add_table(rows=0, cols=len(cleaned_rows[0]))
    for values in cleaned_rows:
        row = new_table.add_row().cells
        for idx, value in enumerate(values):
            row[idx].text = value


def clean_file(input_path: Path, output_path: Path) -> None:
    src = Document(str(input_path))
    dst = Document()
    for block in iter_blocks(src):
        if isinstance(block, Paragraph):
            if should_drop_paragraph(block.text):
                continue
            cleaned = clean_text(block.text)
            if cleaned:
                dst.add_paragraph(cleaned)
        else:
            copy_table(dst, block)
    dst.save(str(output_path))


def main() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    for part in range(1, 16):
        name = f"textbook_part_{part:02d}.docx"
        clean_file(SOURCE_DIR / name, OUTPUT_DIR / name)
    print("Final book ready folder complete.")


if __name__ == "__main__":
    main()
