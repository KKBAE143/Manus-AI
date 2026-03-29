#!/usr/bin/env python3

import os
import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

INPUT_FOLDER = Path("output_book_markers_200_cleaned")
OUTPUT_FOLDER = Path("clean_manuscript_v2")

OUTPUT_FOLDER.mkdir(exist_ok=True)

INLINE_ARTIFACT_PATTERNS = [
    r"\bIf you want,? I'll\b.*",
    r"\bWould you like me to\b.*",
    r"\bWould you like this\b.*",
    r"\bI'll now make\b.*",
    r"\bI'll now draft\b.*",
    r"\bI'll prepare\b.*",
    r"\bI'll generate\b.*",
    r"\bI'll create the Word file\b.*",
    r"\bLet's build\b.*",
    r"\bOnce you reply\b.*",
    r"\bNext module(?: options| choices)?\b.*",
    r"\bNext up, would you like\b.*",
    r"\bSuperb\b.*",
    r"\bAwesome - I'll take it from here\b.*",
    r"\bI can create\b.*",
    r"\bI can prepare\b.*",
    r"\bConfirm these\b.*",
    r"\bJust confirm\b.*",
    r"\bDo you want to include\b.*",
    r"\bShould I include\b.*",
    r"\bBefore I create the Word file\b.*",
    r"\bBefore I finalize\b.*",
    r"\bBefore we proceed\b.*",
    r"\bReady for the next\b.*",
    r"\bStart Disorder\s+\d+\b.*",
    r"\bSay:?\s*Start Disorder\s+\d+\b.*",
    r"\bHere's what I'll do\b.*",
    r"\bNext Step\b.*",
]

INLINE_ARTIFACT_REGEXES = [re.compile(pattern, re.IGNORECASE) for pattern in INLINE_ARTIFACT_PATTERNS]

FRAGMENT_SPLIT_RE = re.compile(
    r"\n+"
    r"|(?<=[.!?])\s+"
    r"|(?=\bSECTION\s+\d+)"
    r"|(?=\bDISORDER\s+\d+)"
    r"|(?=\bMODULE\s+\d+)"
    r"|(?=\b(?:Preamble|Overview|Purpose|Theme|Sample Menu|Case Study|Clinical Insight|Counseling|Dietary Guidelines|Nutritional Goals|Memory Code|Practical Revision Notes|Author's Note|From the Desk|Early Morning|Breakfast|Mid-Morning|Lunch|Snack|Dinner|Bedtime)\b)"
)


def iter_blocks(doc: DocxDocument):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def normalize_fragment(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def is_artifact_fragment(text: str) -> bool:
    candidate = normalize_fragment(text)
    if not candidate:
        return False

    for regex in INLINE_ARTIFACT_REGEXES:
        if regex.search(candidate):
            return True
    return False


def clean_text_block(text: str) -> str:
    raw_fragments = [frag for frag in FRAGMENT_SPLIT_RE.split(text) if frag and frag.strip()]
    kept_fragments = []

    for fragment in raw_fragments:
        cleaned = normalize_fragment(fragment)
        if not cleaned:
            continue
        if is_artifact_fragment(cleaned):
            continue
        kept_fragments.append(cleaned)

    return "\n".join(kept_fragments)


def add_clean_paragraph(dst: Document, text: str) -> None:
    cleaned = clean_text_block(text)
    if cleaned:
        dst.add_paragraph(cleaned)


def add_clean_table(dst: Document, table: Table) -> None:
    cleaned_rows = []
    for row in table.rows:
        cleaned_cells = []
        for cell in row.cells:
            cleaned_cells.append(clean_text_block(cell.text))
        if any(cell.strip() for cell in cleaned_cells):
            cleaned_rows.append(cleaned_cells)

    if not cleaned_rows:
        return

    new_table = dst.add_table(rows=0, cols=len(cleaned_rows[0]))
    for row_values in cleaned_rows:
        row = new_table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].text = value


def clean_file(input_path: Path, output_path: Path) -> None:
    src = Document(str(input_path))
    dst = Document()

    for block in iter_blocks(src):
        if isinstance(block, Paragraph):
            add_clean_paragraph(dst, block.text)
        elif isinstance(block, Table):
            add_clean_table(dst, block)

    dst.save(str(output_path))


def main() -> None:
    for file_name in sorted(os.listdir(INPUT_FOLDER)):
        if file_name.endswith(".docx"):
            clean_file(INPUT_FOLDER / file_name, OUTPUT_FOLDER / file_name)

    print("Sentence-level cleaning complete.")


if __name__ == "__main__":
    main()
