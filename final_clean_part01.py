#!/usr/bin/env python3

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

INPUT_PATH = Path("clean_manuscript_v3/textbook_part_01.docx")
OUTPUT_PATH = Path("textbook_part_01_ultra_clean.docx")

DROP_PARAGRAPH_PATTERNS = [
    re.compile(r"Let's go point by point", re.I),
    re.compile(r"It's totally fine to include your number and email", re.I),
    re.compile(r"If you plan to publish or share publicly online", re.I),
    re.compile(r"Educational Highlight", re.I),
    re.compile(r"Absolutely - I'll include a dynamic index listing", re.I),
    re.compile(r"multi special hospital", re.I),
    re.compile(r"Option 1\s*->\s*Keep your real contacts", re.I),
    re.compile(r"Email:\s*_+", re.I),
    re.compile(r"Phone:\s*_+", re.I),
    re.compile(r"Educational-styled layout", re.I),
    re.compile(r"Should I include both\?", re.I),
    re.compile(r"Age-Sex Coverage", re.I),
    re.compile(r"Right now, the \.docx generation feature is temporarily unstable", re.I),
    re.compile(r"Front & Spine Cover Concept", re.I),
    re.compile(r"Which one do you prefer I do first", re.I),
    re.compile(r"Including a matching back cover", re.I),
    re.compile(r"please confirm these final layout preferences", re.I),
    re.compile(r"I'll integrate these formatting choices", re.I),
    re.compile(r"This makes it look and feel like a published clinical manual", re.I),
    re.compile(r"dynamic index listing", re.I),
    re.compile(r"Final Direction \(Confirmed for Word Manual Build\)", re.I),
    re.compile(r"Memory Code \+ Food Exchange Tables\) in the professional Word layout", re.I),
    re.compile(r"Focus Nutrition Tables, which is like your instant recall system", re.I),
    re.compile(r"Now it's officially ready to go for Word document layout", re.I),
    re.compile(r"The Author's Note in your Word manual", re.I),
    re.compile(r"Add a header & footer on all inner pages", re.I),
    re.compile(r"What I'll Add to the Back Cover When You're Ready", re.I),
    re.compile(r"Here's how we'll finalize the opening section", re.I),
    re.compile(r"This layout is structured to give that premium first impression", re.I),
    re.compile(r"We'll now begin Section 2 - Nutrient Calculation Mastery", re.I),
    re.compile(r"the Section-1 Introduction \(Core Fundamentals\) in this same polished layout", re.I),
    re.compile(r"Before I begin building Section", re.I),
    re.compile(r"Which would you like next\?", re.I),
    re.compile(r"Proceed with Case", re.I),
    re.compile(r"Beautiful - we're heading into", re.I),
    re.compile(r"Let's begin the next premium section", re.I),
    re.compile(r"continuation in the manual's structure", re.I),
    re.compile(r"the next logical and clinically linked module", re.I),
    re.compile(r"your Platinum Manual is now building", re.I),
    re.compile(r"ANSWER TO YOUR CORE QUESTION", re.I),
    re.compile(r"FINAL STATUS \(FOR YOUR PEACE\)", re.I),
    re.compile(r"IMPORTANT: YOU DO NOT NEED TO LIST EVERYTHING", re.I),
]

DROP_LINE_PATTERNS = [
    re.compile(r"^\[SOURCE PAGE \d+\]$", re.I),
    re.compile(r"^and for my$", re.I),
    re.compile(r"^It sounds premium, professional, and clinic-ready\.?$", re.I),
    re.compile(r"^training\)\?$", re.I),
    re.compile(r"^\(e\.g\., Dt\.$", re.I),
    re.compile(r"^\(B\) Educational styled.*$", re.I),
    re.compile(r"^Should I include your email/phone placeholders.*$", re.I),
    re.compile(r"^Should I include both\?$", re.I),
    re.compile(r"^1\.$", re.I),
    re.compile(r"^2\.$", re.I),
    re.compile(r"^3\.$", re.I),
]

DROP_SENTENCE_PATTERNS = [
    re.compile(r"so,?\.*in the section 3 give me more options.*", re.I),
    re.compile(r"more veg options because maximum as everyone can eat veg.*", re.I),
    re.compile(r"provide drills with 3\s*-?4 or 5 examples.*", re.I),
    re.compile(r"even you can add more components.*", re.I),
    re.compile(r"extended tables, cases, and space for your daily notes\)?", re.I),
    re.compile(r"in that mention what is for what.*", re.I),
    re.compile(r"like that if we can do its better na!?", re.I),
    re.compile(r"decoding, and practicing real case-based examples for every section\.?", re.I),
    re.compile(r"what it gives\).*", re.I),
    re.compile(r"the patient also become happy.*", re.I),
    re.compile(r"let'?s train your brain to think like a clinical calculator\.?", re.I),
    re.compile(r"it becomes your \"one-glance reference\".*", re.I),
    re.compile(r"you'?re thinking like a lead clinical nutritionist.*", re.I),
    re.compile(r"absolutely - i'?ll include.*", re.I),
    re.compile(r"it sounds premium, professional, and clinic-ready\.?", re.I),
    re.compile(r"your professional vision statement.*", re.I),
    re.compile(r"which one do you prefer.*", re.I),
    re.compile(r"should i include that too\??", re.I),
    re.compile(r"then here'?s how i'?ll design.*", re.I),
    re.compile(r"i'?ll keep it short, powerful, and purposeful.*", re.I),
    re.compile(r"this page will appear right after the cover.*", re.I),
    re.compile(r"yes include this also.*", re.I),
    re.compile(r"let'?s say a patient'?s plan is.*", re.I),
    re.compile(r"final direction.*word manual build.*", re.I),
    re.compile(r"it will read like a global expert'?s preface.*", re.I),
    re.compile(r"so the tone will be:.*", re.I),
    re.compile(r"memory code \+ food exchange tables\).*", re.I),
    re.compile(r"once that'?s done, we'?ll continue.*", re.I),
    re.compile(r"here'?s how we'?ll start.*", re.I),
    re.compile(r"it'?ll include your big table.*", re.I),
    re.compile(r"focus nutrition tables, which is like your instant recall system.*", re.I),
    re.compile(r"think of this section as your clinical .* nutrition map.*", re.I),
    re.compile(r"i'?ll structure it beautifully.*", re.I),
    re.compile(r"now it'?s officially ready to go for word document layout.*", re.I),
    re.compile(r"apply the same .* style.*", re.I),
    re.compile(r"section 1\) and reformat it with your manual'?s style.*", re.I),
    re.compile(r"here'?s how i'?ll integrate it.*", re.I),
    re.compile(r"you can easily print or laminate.*", re.I),
    re.compile(r"add a header & footer on all inner pages.*", re.I),
    re.compile(r"what i'?ll add to the back cover.*", re.I),
    re.compile(r"if you say yes, i'?ll add.*", re.I),
    re.compile(r"that way, your manual stays future-ready.*", re.I),
    re.compile(r"here'?s how we'?ll finalize the opening section.*", re.I),
    re.compile(r"a formatted front cover layout.*", re.I),
    re.compile(r"this layout is structured to give that premium first impression.*", re.I),
    re.compile(r"we'?ll now begin section 2.*", re.I),
    re.compile(r"it will look and read exactly like your section 1 layout.*", re.I),
    re.compile(r"here'?s what this section will cover.*", re.I),
    re.compile(r"it'?ll include.*", re.I),
    re.compile(r"before i begin building.*", re.I),
    re.compile(r"before i start writing.*", re.I),
    re.compile(r"which would you like next.*", re.I),
    re.compile(r"anything best from your side.*", re.I),
    re.compile(r"anything of your choice.*", re.I),
    re.compile(r"educational layout next.*", re.I),
    re.compile(r"beautiful - we'?re heading into.*", re.I),
    re.compile(r"let'?s begin our next.*", re.I),
    re.compile(r"let'?s begin the next premium section.*", re.I),
    re.compile(r"the next logical and clinically linked module.*", re.I),
    re.compile(r"the next most logical and powerful step.*", re.I),
    re.compile(r"your platinum manual is now building.*", re.I),
    re.compile(r"best practice \(recommended\).*", re.I),
    re.compile(r"you can peacefully proceed.*", re.I),
    re.compile(r"what to do enjoy your green tea.*", re.I),
    re.compile(r"you'?re doing exceptional work.*", re.I),
    re.compile(r"for your clarifying.*", re.I),
    re.compile(r"final status \(for your peace\).*", re.I),
    re.compile(r"you can now stop worrying about structure forever.*", re.I),
    re.compile(r"gpt, this one also for manual.*", re.I),
    re.compile(r"approved correct professional.*", re.I),
    re.compile(r"from this point onward.*belongs to.*", re.I),
    re.compile(r"important: you do not need to list everything.*", re.I),
    re.compile(r"you can simply send one short confirmation message.*", re.I),
]


def iter_blocks(doc: DocxDocumentType):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def norm(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()


def should_drop_paragraph(text: str) -> bool:
    candidate = norm(text)
    if not candidate:
        return False
    for pattern in DROP_PARAGRAPH_PATTERNS:
        if pattern.search(candidate):
            return True
    return False


def clean_text(text: str) -> str:
    lines = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if any(pat.match(line) for pat in DROP_LINE_PATTERNS):
            continue
        lines.append(line)

    cleaned = "\n".join(lines)
    sentences = re.split(r"(?<=[.!?])\s+", cleaned)
    kept = []
    for sentence in sentences:
        candidate = norm(sentence)
        if not candidate:
            continue
        if any(pat.search(candidate) for pat in DROP_SENTENCE_PATTERNS):
            continue
        kept.append(candidate)
    return "\n".join(kept).strip()


def copy_table(dst: DocxDocumentType, table: Table) -> None:
    cleaned_rows = []
    for row in table.rows:
        values = [clean_text(cell.text) for cell in row.cells]
        if any(value.strip() for value in values):
            cleaned_rows.append(values)

    if not cleaned_rows:
        return

    new_table = dst.add_table(rows=0, cols=len(cleaned_rows[0]))
    for values in cleaned_rows:
        row = new_table.add_row().cells
        for idx, value in enumerate(values):
            row[idx].text = value


def main() -> None:
    src = Document(str(INPUT_PATH))
    dst = Document()

    for block in iter_blocks(src):
        if isinstance(block, Paragraph):
            if should_drop_paragraph(block.text):
                continue
            cleaned = clean_text(block.text)
            if cleaned:
                dst.add_paragraph(cleaned)
        elif isinstance(block, Table):
            copy_table(dst, block)

    dst.save(str(OUTPUT_PATH))


if __name__ == "__main__":
    main()
