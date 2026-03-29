#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer

BASE_DIR = Path("book_merge_ready")
OUTPUT_DOCX = Path("Clinical_Nutrition_Master_Textbook.docx")

PARTS = [BASE_DIR / f"textbook_part_{i:02d}.docx" for i in range(1, 16)]
APPENDIX = BASE_DIR / "appendix_reference.docx"


def add_page_break(doc_path: Path) -> Path:
    doc = Document(str(doc_path))
    doc.add_page_break()
    temp_path = doc_path.with_name(doc_path.stem + "_with_break.docx")
    doc.save(str(temp_path))
    return temp_path


def main() -> None:
    master = Document(str(PARTS[0]))
    composer = Composer(master)
    temp_files = []

    for part in PARTS[1:]:
        temp = add_page_break(part)
        temp_files.append(temp)
        composer.append(Document(str(temp)))

    appendix_heading_doc = Document()
    appendix_heading_doc.add_page_break()
    appendix_heading_doc.add_heading("APPENDIX", level=1)
    appendix_heading_doc.add_paragraph("Clinical Nutrition Reference Tables")
    heading_temp = BASE_DIR / "_appendix_heading.docx"
    appendix_heading_doc.save(str(heading_temp))
    temp_files.append(heading_temp)
    composer.append(Document(str(heading_temp)))
    composer.append(Document(str(APPENDIX)))

    composer.save(str(OUTPUT_DOCX))

    for temp in temp_files:
        if temp.exists():
            temp.unlink()


if __name__ == "__main__":
    main()
