#!/usr/bin/env python3

import re
from pathlib import Path

from docx import Document

SOURCE_DIR = Path("final_book_ready")
OUTPUT_DIR = Path("book_merge_ready")
APPENDIX_NAME = "appendix_reference.docx"

STRUCTURE_LINE_RE = re.compile(
    r"\b(?:SECTION|MODULE|PART|FOUNDATION MODULE|THERAPEUTIC MODULE|PRACTICE MODULE|PROFESSIONAL TOOLKIT)\b",
    re.I,
)
APPENDIX_START_RE = re.compile(
    r"\b(?:GLOSSARY|HOUSEHOLD MEASURES|CONVERSION CHART|KCAL EQUIVALENCE|REFERENCE INDEX)\b",
    re.I,
)
VALIDATION_RE = re.compile(
    r"\b(?:SECTION|MODULE|TOOLKIT|REFERENCE BUILD|CONVERSION CHART)\b",
    re.I,
)
ORPHAN_NUMBER_RE = re.compile(r"^\d+[.)]?$")


def normalize_paragraph(text: str) -> str:
    lines = []
    for raw_line in text.splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        if not line:
            if lines and lines[-1] == "":
                continue
            lines.append("")
            continue
        if ORPHAN_NUMBER_RE.match(line):
            continue
        if STRUCTURE_LINE_RE.search(line):
            continue
        lines.append(line)

    while lines and lines[0] == "":
        lines.pop(0)
    while lines and lines[-1] == "":
        lines.pop()

    cleaned = "\n".join(lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    return cleaned


def main() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    appendix_doc = Document()
    appendix_has_content = False

    for part in range(1, 16):
        name = f"textbook_part_{part:02d}.docx"
        src = Document(str(SOURCE_DIR / name))
        dst = Document()
        in_appendix = False

        for para in src.paragraphs:
            raw = para.text or ""
            if APPENDIX_START_RE.search(raw):
                in_appendix = True

            cleaned = normalize_paragraph(raw)
            if not cleaned:
                continue

            target = appendix_doc if in_appendix else dst
            target.add_paragraph(cleaned)
            if in_appendix:
                appendix_has_content = True

        for table in src.tables:
            # Preserve table placement in the main part only; appendix extraction is paragraph-triggered.
            new_table = dst.add_table(rows=0, cols=len(table.columns))
            for row in table.rows:
                values = []
                for cell in row.cells:
                    cleaned = normalize_paragraph(cell.text)
                    values.append(cleaned)
                if not any(values):
                    continue
                new_row = new_table.add_row().cells
                for idx, value in enumerate(values):
                    new_row[idx].text = value

        dst.save(str(OUTPUT_DIR / name))

    if not appendix_has_content:
        appendix_doc.add_paragraph("")
    appendix_doc.save(str(OUTPUT_DIR / APPENDIX_NAME))
    print("Book merge folder ready.")


if __name__ == "__main__":
    main()
