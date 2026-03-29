#!/usr/bin/env python3

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

INPUT_FOLDER = Path("output_book_markers_200_cleaned")
OUTPUT_FOLDER = Path("clean_manuscript_v3")
TARGET_FILES = [f"textbook_part_{i:02d}.docx" for i in range(1, 16)]

PART01_SENTENCE_DROP_PATTERNS = [
    re.compile(r"\badd another section\b", re.I),
    re.compile(r"\bwe can structure\b", re.I),
    re.compile(r"\blet'?s design\b", re.I),
    re.compile(r"\bonce we finalize\b", re.I),
    re.compile(r"\bfollowing your design\b", re.I),
    re.compile(r"\bi will generate\b", re.I),
    re.compile(r"\bi will create\b", re.I),
    re.compile(r"\bi will prepare\b", re.I),
    re.compile(r"\bbefore i generate\b", re.I),
    re.compile(r"\bwould you like\b", re.I),
    re.compile(r"\bonce you confirm\b", re.I),
    re.compile(r"\blayout option\b", re.I),
    re.compile(r"\bformat choice\b", re.I),
    re.compile(r"\bcover page\b", re.I),
    re.compile(r"\bfront page\b", re.I),
    re.compile(r"\btitle page\b", re.I),
    re.compile(r"\bemail placeholder\b", re.I),
    re.compile(r"\bphone placeholder\b", re.I),
    re.compile(r"\bclinic branding\b", re.I),
    re.compile(r"\bemail\s*:\s*\S+@\S+", re.I),
    re.compile(r"\bph\.?no\s*[:;]\s*\+?\d+", re.I),
    re.compile(r"\bphone\s*[:;]\s*\+?\d+", re.I),
    re.compile(r"\bclinical nutrition practice accelerator\b", re.I),
    re.compile(r"\bmanual structure\b", re.I),
    re.compile(r"\bplanned structure\b", re.I),
    re.compile(r"\bsection\s+\w+\s+will contain\b", re.I),
    re.compile(r"\bi read the manual\b", re.I),
    re.compile(r"\bfrom my end\b", re.I),
    re.compile(r"\blet'?s design this upgrade\b", re.I),
]

OUTPUT_FOLDER.mkdir(exist_ok=True)

INLINE_ARTIFACT_PATTERNS = [
    r"\byess please!?\b.*",
    r"\bbuddy!?\b.*",
    r"\bA BIG YESS!?\b.*",
    r"\bWould you like\b.*",
    r"\bIf you want\b.*",
    r"\bI'll now\b.*",
    r"\bI will now\b.*",
    r"\bI'll prepare\b.*",
    r"\bI'll generate\b.*",
    r"\bI'll create\b.*Word file.*",
    r"\bLet's build\b.*",
    r"\bOnce you reply\b.*",
    r"\bOnce you pick\b.*",
    r"\bWhich do you prefer\b.*",
    r"\bHere'?s what we'll do step-by-step\b.*",
    r"\bconfirm this structure before I create\b.*",
    r"\bIf this final structure is confirmed\b.*",
    r"\bThat seals it beautifully\b.*",
    r"\bSubstitution Charts next\?.*",
    r"\bHow to Use These Templates\b.*",
    r"\bfull examples\) next\?.*",
    r"\bFast Choice Builder\) next\?.*",
    r"\bTables next\?.*",
    r"\bCase Study .* next\?.*",
    r"\bReady for the next\b.*",
    r"\bStart Disorder\b.*",
    r"\bSure buddy\b.*",
    r"\bProceed to Disorder\s+\d+\b.*",
    r"\bJust say\s+[\"']?Proceed[\"']?\b.*",
    r"\bSay\s+[\"']?Proceed[\"']?\s+when ready\b.*",
    r"\bContinuing in the same Platinum Manual\b.*",
    r"\bContinuing in the same Platinum Manual \|.*",
    r"\bContinuing seamlessly\b.*",
    r"\bReady when you are\b.*",
    r"\bWe now continue\b.*",
    r"\bWe now move into\b.*",
    r"\bThis module includes\b.*",
    r"\bsame premium quality\b.*",
    r"\bWord export feature\b.*",
    r"\bfinalize this\b.*",
    r"\bfull manual\b.*Word.*",
    r"\bbuild the full Word version\b.*",
    r"\bit'?ll look like\b.*",
    r"\bonce i create\b.*",
    r"\bonce we create\b.*",
    r"\bi'?ll immediately generate\b.*",
    r"\bthis section will appear\b.*",
    r"\bbelow is the continuation\b.*",
    r"\bfinalized structure\b.*",
    r"\bplanned structure inside the file\b.*",
    r"\bfinal structure\b.*",
    r"\bi read the manual\b.*",
    r"\bfrom my end i want\b.*",
    r"\byes and finally\b.*",
    r"\bthat'?s exactly the mindset\b.*",
    r"\blet'?s design this upgrade\b.*",
    r"\bthis will be your professional toolbook\b.*",
    r"\bpersonal clinic-ready\b.*",
    r"\bso your Word file\b.*",
    r"\bfor title i like\b.*",
    r"\bwhat clinic name\b.*",
    r"\bword file\b.*",
    r"\bgenerate the word\b.*",
    r"\bexport\b.*",
    r"\bcompile\b.*",
    r"\blayout selection\b.*",
    r"\bformatting choice\b.*",
    r"\btitle page\b.*",
    r"\bfront page\b.*",
    r"\bcover page\b.*",
    r"\bemail placeholder\b.*",
    r"\bphone placeholder\b.*",
    r"\bclinic branding\b.*",
    r"\btagline\b.*",
    r"\blayout preference\b.*",
    r"\bplanned structure inside the file\b.*",
    r"\bsection\s+\d+\s+will contain\b.*",
    r"\blayout selection option\b.*",
    r"\bhere'?s the plan\b.*",
    r"\bhere'?s the start of your next chapter draft\b.*",
    r"\bthis is the section that makes your counseling lightning-fast\b.*",
    r"\bthis section will make you\b.*",
    r"\bclinical nutrition matrix\b.*one-glance reference\b.*",
    r"\bcover page\b.*",
    r"\bfront page\b.*",
    r"\btitle page\b.*",
    r"\beducational highlight layout\b.*",
    r"\bprofessional minimal\b.*",
    r"\blayout choice\b.*",
    r"\bclinic name\b.*",
    r"\bprofessional name/title\b.*",
    r"\bformat plan\b.*",
    r"\bplacement order\b.*",
    r"\bYes gpt we can proceed\b.*",
    r"\bfor my satisfaction\b.*",
    r"\bYou're absolutely right to pause and cross-check\b.*",
    r"\bLet me give you a clear, calm, one-screen confirmation\b.*",
    r"\bSO, TO YOUR QUESTION\b.*",
    r"\bANSWER: YES - PLEASE SEND IT\b.*",
    r"\bJust list what each tab roughly contains\b.*",
    r"\bOnce you send that\b.*",
    r"\bWhenever you're ready- Send the tab list\b.*",
    r"\bI'm aligned with you\b.*",
    r"\bChalo gpt\b.*",
    r"\bOkay i understand it now\b.*",
    r"\bExactly - you've understood it 100% correctly\b.*",
    r"\bYou don't need to remember\b.*",
    r"\bI'll mentally index everything\b.*",
    r"\bFuture content will auto-align\b.*",
    r"\bCompilation becomes clean Word file generation becomes\b.*",
    r"\bDisorder\s+\d+\s+COMPLETE\b.*",
    r"\bDISORDER\s+\d+\s+COMPLETE\b.*",
    r"^NEXT$",
    r"\bNEXT\b.*",
]

DROP_EXACT = {
    "page?",
    "3.",
    "next",
}

INLINE_ARTIFACT_REGEXES = [re.compile(pattern, re.IGNORECASE) for pattern in INLINE_ARTIFACT_PATTERNS]

PERSONAL_DATA_RE = [
    re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}"),
    re.compile(r"(?:ph\.?\s*(?:no\.?|num\.?|number)?\s*[:;]?\s*|phone\s*[:;]\s*|mobile\s*[:;]\s*|contact\s*[:;]\s*)[\+]?\d[\d\s\-]{7,14}", re.I),
    re.compile(r"(?<!\d)\d{10}(?!\d)"),
]

CHAT_DROP_PATTERNS_V3 = [
    re.compile(r"\bi read the manual\b", re.I),
    re.compile(r"\bfrom my end i want\b", re.I),
    re.compile(r"\bas of now i work\b", re.I),
    re.compile(r"\bsoon gonna leave\b", re.I),
    re.compile(r"\bin that mention what is for what\b", re.I),
    re.compile(r"\bthat'?s exactly the mindset\b", re.I),
    re.compile(r"\byou'?re thinking like a lead clinical\b", re.I),
    re.compile(r"\blet'?s design this upgrade\b", re.I),
    re.compile(r"\bgreat choice[!.]?\s*starting with\b", re.I),
    re.compile(r"\bnext,\s*would you like me to create\b", re.I),
    re.compile(r"\bwould you like me to create section\b", re.I),
    re.compile(r"\bthis next section will be your secret\b", re.I),
    re.compile(r"\bif you want,?\s*i'?ll now make you\b", re.I),
    re.compile(r"\byess please[!.]?\b", re.I),
    re.compile(r"\bspot on[!.]?\b", re.I),
    re.compile(r"\b-\s*let'?s roll[!.]?\b", re.I),
    re.compile(r"\b-\s*great choice[!.]?\b", re.I),
    re.compile(r"\bkaladhar\s*[-–]\s*we'?re now building\b", re.I),
    re.compile(r"\bit'?ll make you unbeatable\b", re.I),
    re.compile(r"\byour clients will never need a second opinion\b", re.I),
    re.compile(r"\bonce we create this,? you'?ll have\b", re.I),
    re.compile(r"\bword\s*\+\s*pdf version\b", re.I),
    re.compile(r"\bwhich do you prefer\b", re.I),
    re.compile(r"\bconfirm these\s*\d*\s*final details\b", re.I),
    re.compile(r"\bfinal structure[:\s]*['\"]?platinum\b", re.I),
    re.compile(r"\bplanned structure inside the file\b", re.I),
    re.compile(r"\bfront page setup\b.*edition\b", re.I),
    re.compile(r"\blayout selection option\b", re.I),
    re.compile(r"\bformat plan\b", re.I),
    re.compile(r"\bplacement order\b", re.I),
    re.compile(r"\bthis will be your professional toolbook\b", re.I),
    re.compile(r"\bmaster dietitian daily clinical bible\b", re.I),
    re.compile(r"\bdiet modification master manual\b", re.I),
    re.compile(r"\bnavath\s*kaladhar\b", re.I),
    re.compile(r"\bnavathkaladhar\b", re.I),
]

PIPE_TABLE_RE_V3 = re.compile(r"^\s*\|.+\|.+\|\s*$")

NUMBER_EMOJI_MAP_V3 = {
    "1\u20e3": "1.", "2\u20e3": "2.", "3\u20e3": "3.", "4\u20e3": "4.",
    "5\u20e3": "5.", "6\u20e3": "6.", "7\u20e3": "7.", "8\u20e3": "8.",
    "9\u20e3": "9.", "\U0001f51f": "10.",
}
NUMBER_EMOJI_RE_V3 = re.compile("|".join(re.escape(k) for k in NUMBER_EMOJI_MAP_V3))

TEXT_FUSION_RE_V3 = re.compile(r"([a-z])([A-Z])")
DIGIT_LETTER_RE_V3 = re.compile(r"(\d)([A-Za-z])")
LETTER_DIGIT_RE_V3 = re.compile(r"([a-z])(\d)")

KNOWN_TABLE_HEADER_RE_V3 = [
    re.compile(r"condition\s*/\s*key nutrition focus\s*/\s*supportive nutrients", re.I),
    re.compile(r"foods to include\s*/\s*foods to avoid\s*/\s*clinical tips", re.I),
]

FRAGMENT_SPLIT_RE = re.compile(
    r"\n+"
    r"|(?<=[.!?])\s+"
    r"|(?=\bSECTION\s+\d+)"
    r"|(?=\bDISORDER\s+\d+)"
    r"|(?=\bMODULE\s+\d+)"
    r"|(?=\bNEXT\b)"
    r"|(?=\bProceed to Disorder\s+\d+)"
    r"|(?=\bSay\s+[\"']?Proceed[\"']?)"
    r"|(?=\bJust say\s+[\"']?Proceed[\"']?)"
    r"|(?=\b(?:Preamble|Overview|Purpose|Theme|Sample Menu|Case Study|Clinical Insight|Counseling|Dietary Guidelines|Nutritional Goals|Memory Code|Practical Revision Notes|Author's Note|From the Desk|Early Morning|Breakfast|Mid-Morning|Lunch|Snack|Dinner|Bedtime|INDEX|COVER PAGE|Title|Final Design|FOLLOW-UP|INTRODUCTION)\b)"
)


def iter_blocks(doc: DocxDocumentType):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _replace_number_emoji_v3(m: "re.Match") -> str:
    return NUMBER_EMOJI_MAP_V3.get(m.group(0), m.group(0))


def _strip_personal_data(text: str) -> str:
    for pat in PERSONAL_DATA_RE:
        text = pat.sub("", text)
    text = re.sub(r"[ \t]{2,}", " ", text).strip()
    return text


def _fix_text_fusion(text: str) -> str:
    text = TEXT_FUSION_RE_V3.sub(r"\1 \2", text)
    text = DIGIT_LETTER_RE_V3.sub(r"\1 \2", text)
    text = LETTER_DIGIT_RE_V3.sub(r"\1 \2", text)
    return text


_STANDALONE_NUMBER_RE_V3 = re.compile(r"^\d{1,2}\.\s*$")


def normalize_fragment(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = NUMBER_EMOJI_RE_V3.sub(_replace_number_emoji_v3, text)
    text = _strip_personal_data(text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def is_artifact_fragment(text: str) -> bool:
    candidate = normalize_fragment(text)
    if not candidate:
        return False
    if candidate.lower() in DROP_EXACT:
        return True

    if _STANDALONE_NUMBER_RE_V3.match(candidate):
        return True

    if PIPE_TABLE_RE_V3.match(candidate):
        return True

    for regex in INLINE_ARTIFACT_REGEXES:
        if regex.search(candidate):
            return True

    for pat in CHAT_DROP_PATTERNS_V3:
        if pat.search(candidate):
            return True

    return False


def should_drop_part01_sentence(text: str) -> bool:
    candidate = normalize_fragment(text)
    if not candidate:
        return False
    for pattern in PART01_SENTENCE_DROP_PATTERNS:
        if pattern.search(candidate):
            return True
    return False


def clean_text_block(text: str, is_part01: bool = False) -> str:
    raw_fragments = [frag for frag in FRAGMENT_SPLIT_RE.split(text) if frag and frag.strip()]
    kept_fragments = []

    for fragment in raw_fragments:
        cleaned = normalize_fragment(fragment)
        if not cleaned or is_artifact_fragment(cleaned):
            continue
        if is_part01:
            sentence_bits = [bit for bit in re.split(r"(?<=[.!?])\s+", cleaned) if bit.strip()]
            sentence_bits = [bit for bit in sentence_bits if not should_drop_part01_sentence(bit)]
            cleaned = " ".join(sentence_bits).strip()
            if not cleaned:
                continue
        cleaned = _fix_text_fusion(cleaned)
        kept_fragments.append(cleaned)

    return "\n".join(kept_fragments)


def add_clean_paragraph(dst: DocxDocumentType, text: str, is_part01: bool = False) -> None:
    cleaned = clean_text_block(text, is_part01=is_part01)
    if cleaned:
        dst.add_paragraph(cleaned)


def _is_table_header_row(cells: list) -> bool:
    joined = " / ".join(c for c in cells if c)
    for pat in KNOWN_TABLE_HEADER_RE_V3:
        if pat.search(joined):
            return True
    return False


def add_clean_table(dst: DocxDocumentType, table: Table, is_part01: bool = False, seen_headers: set = None) -> None:
    if seen_headers is None:
        seen_headers = set()
    cleaned_rows = []
    for row in table.rows:
        cleaned_cells = [clean_text_block(cell.text, is_part01=is_part01) for cell in row.cells]
        if not any(cell.strip() for cell in cleaned_cells):
            continue
        if _is_table_header_row(cleaned_cells):
            key = re.sub(r"\s+", " ", " / ".join(c for c in cleaned_cells if c).lower().strip())
            if key in seen_headers:
                continue
            seen_headers.add(key)
        cleaned_rows.append(cleaned_cells)

    if not cleaned_rows:
        return

    new_table = dst.add_table(rows=0, cols=len(cleaned_rows[0]))
    for row_values in cleaned_rows:
        row = new_table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].text = value


def clean_file(input_path: Path, output_path: Path) -> None:
    src = Document(str(input_path))
    dst = Document()
    is_part01 = input_path.name == "textbook_part_01.docx"
    seen_headers: set = set()

    for block in iter_blocks(src):
        if isinstance(block, Paragraph):
            add_clean_paragraph(dst, block.text, is_part01=is_part01)
        elif isinstance(block, Table):
            add_clean_table(dst, block, is_part01=is_part01, seen_headers=seen_headers)

    dst.save(str(output_path))


def main() -> None:
    OUTPUT_FOLDER.mkdir(exist_ok=True)
    processed = 0
    for file_name in TARGET_FILES:
        input_path = INPUT_FOLDER / file_name
        if not input_path.exists():
            print(f"  SKIP (not found): {file_name}")
            continue
        output_path = OUTPUT_FOLDER / file_name
        print(f"  Cleaning {file_name}...")
        clean_file(input_path, output_path)
        processed += 1
    print(f"Surgical v3 cleaning complete. Processed {processed} file(s).")


if __name__ == "__main__":
    main()
