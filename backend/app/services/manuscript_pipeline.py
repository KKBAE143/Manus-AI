import json
import os
import sys
import threading
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

import fitz
from docx import Document as DocxDocument
from docxcompose.composer import Composer
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.document import (
    Artifact,
    ArtifactType,
    Chunk,
    ChunkStageRun,
    ChunkStageStatus,
    ChunkStatus,
    Document,
    DocumentConfig,
    DocumentStatus,
    ExportPart,
    Job,
    JobStatus,
    MergeJob,
    MergeStatus,
    PartStatus,
    ProjectEventLog,
    SplitMode,
)

ROOT_DIR = Path(settings.PROJECT_ROOT)
if str(ROOT_DIR) not in sys.path:
    sys.path.append(str(ROOT_DIR))

from pdf_to_clean_docx import (  # noqa: E402
    PageRecord,
    build_docx_from_blocks,
    build_docx_part,
    clean_page_text,
    content_items_to_dict,
    content_items_to_plain_text,
    extract_page_rich,
    write_manifest,
)


_JOB_CONFIG_LOCK = threading.Lock()
_JOB_CONFIG_CACHE: Dict[str, Dict[str, Any]] = {}


def _cache_document_config(document_id: str, config: "DocumentConfig") -> None:
    with _JOB_CONFIG_LOCK:
        _JOB_CONFIG_CACHE[document_id] = {k: v for k, v in vars(config).items() if not k.startswith("_")}


def _get_cached_document_config(document_id: str, db: Session) -> "DocumentConfig":
    with _JOB_CONFIG_LOCK:
        cached = _JOB_CONFIG_CACHE.get(document_id)
    if cached is not None:
        obj = DocumentConfig.__new__(DocumentConfig)
        for k, v in cached.items():
            setattr(obj, k, v)
        return obj
    result = db.query(DocumentConfig).filter(DocumentConfig.document_id == document_id).first()
    if result is not None:
        _cache_document_config(document_id, result)
    return result


def _evict_document_config(document_id: str) -> None:
    with _JOB_CONFIG_LOCK:
        _JOB_CONFIG_CACHE.pop(document_id, None)


STAGES = [
    (1, "queued", "Queued"),
    (2, "inspect", "Inspecting PDF"),
    (3, "plan", "Planning Chunks"),
    (4, "extract", "Extracting Pages"),
    (5, "clean_pass_1", "Cleanup Pass 1"),
    (6, "clean_pass_2", "Cleanup Pass 2"),
    (7, "final_normalize", "Final Normalize"),
    (8, "ai_transform", "AI Transformation"),
    (9, "part_generate", "Generating DOCX Parts"),
    (10, "appendix_extract", "Appendix Extraction"),
    (11, "merge_prep", "Merge Preparation"),
    (12, "completed", "Completed"),
]


def get_project_dir(document_id: str) -> Path:
    project_dir = Path(settings.PROJECTS_DIR) / document_id
    project_dir.mkdir(parents=True, exist_ok=True)
    return project_dir


def get_outputs_dir(document_id: str) -> Path:
    output_dir = get_project_dir(document_id) / "parts"
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def get_chunks_dir(document_id: str) -> Path:
    path = get_project_dir(document_id) / "chunks"
    path.mkdir(parents=True, exist_ok=True)
    return path


def get_manifests_dir(document_id: str) -> Path:
    path = get_project_dir(document_id) / "manifests"
    path.mkdir(parents=True, exist_ok=True)
    return path


def get_logs_dir(document_id: str) -> Path:
    path = get_project_dir(document_id) / "logs"
    path.mkdir(parents=True, exist_ok=True)
    return path


def get_merge_dir(document_id: str) -> Path:
    path = get_project_dir(document_id) / "merge"
    path.mkdir(parents=True, exist_ok=True)
    return path


def get_appendix_path(document_id: str) -> Path:
    return get_merge_dir(document_id) / "appendix_reference.docx"


_LOG_BATCH_SIZE = 10
_log_event_counter: dict[int, int] = {}


def log_event(db: Session, document_id: str, message: str, *, level: str = "INFO", event_type: str = "pipeline", stage_key: Optional[str] = None, chunk_id: Optional[str] = None, flush: bool = False) -> None:
    db.add(
        ProjectEventLog(
            document_id=document_id,
            level=level,
            event_type=event_type,
            message=message,
            stage_key=stage_key,
            chunk_id=chunk_id,
        )
    )
    db_id = id(db)
    count = _log_event_counter.get(db_id, 0) + 1
    _log_event_counter[db_id] = count
    if flush or count % _LOG_BATCH_SIZE == 0:
        db.commit()
        if flush:
            _log_event_counter.pop(db_id, None)


def flush_log_events(db: Session) -> None:
    db_id = id(db)
    if _log_event_counter.get(db_id, 0) > 0:
        db.commit()
    _log_event_counter.pop(db_id, None)


def register_artifact(db: Session, document_id: str, artifact_type: ArtifactType, label: str, path: Path, chunk_id: Optional[str] = None) -> Artifact:
    artifact = Artifact(
        document_id=document_id,
        chunk_id=chunk_id,
        artifact_type=artifact_type,
        label=label,
        path=str(path),
        size_bytes=path.stat().st_size if path.exists() else None,
    )
    db.add(artifact)
    db.commit()
    db.refresh(artifact)
    return artifact


def create_chunk_stage_run(db: Session, chunk: Chunk, stage_key: str, stage_name: str, *, status: ChunkStageStatus = ChunkStageStatus.processing, message: Optional[str] = None) -> ChunkStageRun:
    stage_run = ChunkStageRun(
        chunk_id=chunk.id,
        stage_key=stage_key,
        stage_name=stage_name,
        status=status,
        message=message,
        retry_count=chunk.retry_count or 0,
        started_at=datetime.utcnow(),
    )
    db.add(stage_run)
    db.commit()
    db.refresh(stage_run)
    return stage_run


def finish_chunk_stage_run(db: Session, stage_run: ChunkStageRun, *, status: ChunkStageStatus, message: Optional[str] = None) -> None:
    stage_run.status = status
    stage_run.message = message or stage_run.message
    stage_run.finished_at = datetime.utcnow()
    db.commit()


def inspect_pdf(file_path: str) -> Dict[str, Any]:
    pdf = fitz.open(file_path)
    try:
        metadata = pdf.metadata or {}
        return {
            "page_count": len(pdf),
            "title": metadata.get("title") or "",
            "author": metadata.get("author") or "",
            "subject": metadata.get("subject") or "",
        }
    finally:
        pdf.close()


def update_job(
    db: Session,
    job: Job,
    *,
    stage: int,
    stage_key: str,
    stage_name: str,
    progress: float,
    message: Optional[str] = None,
    status: Optional[JobStatus] = None,
) -> None:
    job.stage = stage
    job.stage_key = stage_key
    job.stage_name = stage_name
    job.progress_percent = max(0.0, min(progress, 100.0))
    job.progress_message = message
    if status is not None:
        job.status = status
    db.commit()


def run_pipeline(document_id: str, job_id: str, db: Session) -> Dict[str, Any]:
    _evict_document_config(document_id)
    document = db.query(Document).filter(Document.id == document_id).first()
    job = db.query(Job).filter(Job.id == job_id).first()
    config = db.query(DocumentConfig).filter(DocumentConfig.document_id == document_id).first()

    if not document or not job or not config:
        raise ValueError("Document, job, or config not found")
    _cache_document_config(document_id, config)

    project_dir = get_project_dir(document_id)
    outputs_dir = get_outputs_dir(document_id)
    chunks_dir = get_chunks_dir(document_id)
    manifests_dir = get_manifests_dir(document_id)
    logs_dir = get_logs_dir(document_id)
    get_merge_dir(document_id)
    document.storage_root = str(project_dir)
    document.status = DocumentStatus.processing
    job.status = JobStatus.in_progress
    db.commit()
    log_event(db, document_id, f"Job {job_id} started", stage_key="queued")

    update_job(db, job, stage=2, stage_key="inspect", stage_name="Inspecting PDF", progress=5.0, message="Reading PDF metadata")

    pdf = fitz.open(document.local_storage_path)
    total_pages = len(pdf)
    document.page_count = total_pages
    document.status = DocumentStatus.inspected
    db.commit()
    log_event(db, document_id, f"PDF inspected: {total_pages} pages", stage_key="inspect", flush=True)

    start_page = max(1, config.start_page)
    end_page = min(total_pages, config.end_page)

    update_job(db, job, stage=3, stage_key="plan", stage_name="Planning Chunks", progress=10.0, message=f"Planning chunks for pages {start_page}-{end_page}")

    db.query(ExportPart).filter(ExportPart.document_id == document_id).delete(synchronize_session=False)
    db.query(Chunk).filter(Chunk.document_id == document_id).delete(synchronize_session=False)
    db.query(Artifact).filter(Artifact.document_id == document_id).delete(synchronize_session=False)
    db.query(ProjectEventLog).filter(ProjectEventLog.document_id == document_id).delete(synchronize_session=False)
    db.commit()
    log_event(db, document_id, "Chunk planning started", stage_key="plan")

    chunk_ranges = []
    if config.split_mode in (SplitMode.chapters, SplitMode.hybrid):
        from app.services.chapter_detector import check_table_safety, detect_chapter_boundaries_with_fallback, group_chapters_into_chunks
        from app.core.ai_pool import get_ai_pool as _get_ai_pool
        try:
            _plan_pool = _get_ai_pool() if _get_ai_pool().available() else None
        except Exception:
            _plan_pool = None

        boundaries = detect_chapter_boundaries_with_fallback(pdf, start_page, end_page, pool=_plan_pool)
        chapter_boundary_pages = {b.page for b in boundaries}
        chapter_groups = group_chapters_into_chunks(boundaries, start_page, end_page, config.pages_per_docx)
        split_points = [grp[0] for grp in chapter_groups[1:]]
        intra_splits = [sp for sp in split_points if sp not in chapter_boundary_pages]
        adjusted_intra = check_table_safety(pdf, intra_splits, chapter_boundary_pages)
        intra_map = dict(zip(intra_splits, adjusted_intra))
        adjusted_splits = [intra_map.get(sp, sp) for sp in split_points]
        if len(adjusted_splits) == len(chapter_groups) - 1:
            rebuilt = []
            for i, (cs, ce, ct) in enumerate(chapter_groups):
                if i == 0:
                    new_start = cs
                else:
                    new_start = adjusted_splits[i - 1]
                if i < len(chapter_groups) - 1:
                    new_end = adjusted_splits[i] - 1
                else:
                    new_end = ce
                new_start = max(new_start, start_page)
                new_end = min(new_end, end_page)
                if new_start <= new_end:
                    rebuilt.append((new_start, new_end, ct))
            chapter_groups = rebuilt
        chunk_index = 1
        for ch_start, ch_end, ch_title in chapter_groups:
            if ch_start > ch_end or ch_start > end_page:
                continue
            chunk = Chunk(
                document_id=document_id,
                chunk_index=chunk_index,
                page_start=ch_start,
                page_end=ch_end,
                page_count=ch_end - ch_start + 1,
                status=ChunkStatus.queued,
                current_stage="planned",
                chapter_title=ch_title or None,
            )
            db.add(chunk)
            db.commit()
            db.refresh(chunk)
            chunk_ranges.append(chunk)
            chunk_index += 1
    else:
        page_cursor = start_page
        chunk_index = 1
        while page_cursor <= end_page:
            chunk_end = min(end_page, page_cursor + config.pages_per_docx - 1)
            chunk = Chunk(
                document_id=document_id,
                chunk_index=chunk_index,
                page_start=page_cursor,
                page_end=chunk_end,
                page_count=chunk_end - page_cursor + 1,
                status=ChunkStatus.queued,
                current_stage="planned",
            )
            db.add(chunk)
            db.commit()
            db.refresh(chunk)
            chunk_ranges.append(chunk)
            chunk_index += 1
            page_cursor = chunk_end + 1

    document.status = DocumentStatus.planned
    db.commit()
    log_event(db, document_id, f"Planned {len(chunk_ranges)} chunks", stage_key="plan", flush=True)

    page_records = []
    created_parts = []

    total_to_process = max(1, end_page - start_page + 1)

    for chunk in chunk_ranges:
        chunk.status = ChunkStatus.processing
        chunk.current_stage = "extract"
        chunk.started_at = datetime.utcnow()
        db.commit()
        log_event(db, document_id, f"Processing chunk {chunk.chunk_index} ({chunk.page_start}-{chunk.page_end})", stage_key="extract", chunk_id=chunk.id)
        extract_run = create_chunk_stage_run(db, chunk, "extract", "Extract Pages", message="Reading PDF pages")

        raw_lines = []
        cleaned_pages = []
        rich_pages = []
        chunk_table_headers: set = set()
        for source_page in range(chunk.page_start, chunk.page_end + 1):
            page_obj = pdf[source_page - 1]
            try:
                items = extract_page_rich(page_obj)
                plain = content_items_to_plain_text(items)
                raw_text = plain if plain.strip() else page_obj.get_text("text")
            except Exception:
                items = []
                plain = ""
                raw_text = page_obj.get_text("text")
            raw_lines.append(f"===== PAGE {source_page} =====\n{raw_text}")
            cleaned = clean_page_text(raw_text, seen_table_headers=chunk_table_headers)
            keep = bool(plain.strip()) or bool(cleaned.strip())
            page_records.append(PageRecord(source_page=source_page, cleaned_char_count=len(plain or cleaned), kept=keep))
            if keep:
                cleaned_pages.append((source_page, cleaned))
                rich_pages.append((source_page, items))

        raw_path = chunks_dir / f"chunk_{chunk.chunk_index:04d}_raw.txt"
        raw_path.write_text("\n\n".join(raw_lines), encoding="utf-8")
        chunk.raw_text_path = str(raw_path)
        register_artifact(db, document_id, ArtifactType.chunk_text, f"Chunk {chunk.chunk_index} raw text", raw_path, chunk.id)
        finish_chunk_stage_run(db, extract_run, status=ChunkStageStatus.completed, message="Raw extraction complete")

        update_job(
            db,
            job,
            stage=5,
            stage_key="clean_pass_1",
            stage_name="Cleanup Pass 1",
            progress=15.0 + ((chunk.chunk_index - 1) / max(1, len(chunk_ranges))) * 60.0,
            message=f"Cleanup pass 1 for chunk {chunk.chunk_index}/{len(chunk_ranges)}",
        )
        chunk.current_stage = "clean_pass_1"
        db.commit()
        clean1_run = create_chunk_stage_run(db, chunk, "clean_pass_1", "Cleanup Pass 1", message="Applying deterministic cleanup rules")

        cleaned_text = "\n\n".join(f"===== PAGE {page} =====\n{text}" for page, text in cleaned_pages)
        cleaned_path = chunks_dir / f"chunk_{chunk.chunk_index:04d}_cleaned.txt"
        cleaned_path.write_text(cleaned_text, encoding="utf-8")
        chunk.cleaned_text_path = str(cleaned_path)
        register_artifact(db, document_id, ArtifactType.cleaned_text, f"Chunk {chunk.chunk_index} cleaned text", cleaned_path, chunk.id)

        rich_data = [
            {"page": page, "items": content_items_to_dict(items)}
            for page, items in rich_pages
        ]
        rich_path = chunks_dir / f"chunk_{chunk.chunk_index:04d}_rich.json"
        rich_path.write_text(json.dumps(rich_data, ensure_ascii=False, indent=2), encoding="utf-8")
        register_artifact(db, document_id, ArtifactType.cleaned_text, f"Chunk {chunk.chunk_index} rich content JSON", rich_path, chunk.id)

        finish_chunk_stage_run(db, clean1_run, status=ChunkStageStatus.completed, message="Chunk text cleaned")

        chunk.current_stage = "clean_pass_2"
        db.commit()
        clean2_run = create_chunk_stage_run(db, chunk, "clean_pass_2", "Cleanup Pass 2", message="Second cleanup checkpoint")
        finish_chunk_stage_run(db, clean2_run, status=ChunkStageStatus.completed, message="Second cleanup pass complete")

        chunk.current_stage = "final_normalize"
        db.commit()
        normalize_run = create_chunk_stage_run(db, chunk, "final_normalize", "Final Normalize", message="Normalizing chunk output")
        finish_chunk_stage_run(db, normalize_run, status=ChunkStageStatus.completed, message="Normalization complete")

        update_job(
            db,
            job,
            stage=8,
            stage_key="part_generate",
            stage_name="Generating DOCX Parts",
            progress=25.0 + (chunk.chunk_index / max(1, len(chunk_ranges))) * 60.0,
            message=f"Generating DOCX part for chunk {chunk.chunk_index}/{len(chunk_ranges)}",
        )
        chunk.current_stage = "part_generate"
        db.commit()
        part_run = create_chunk_stage_run(db, chunk, "part_generate", "Generate DOCX Part", message="Building DOCX for chunk")

        export_part = _flush_part(
            db,
            document,
            job,
            outputs_dir,
            rich_pages if rich_pages else cleaned_pages,
            chunk.chunk_index,
            config.keep_page_markers,
        )
        created_parts.append(export_part)
        _assign_part_numbers(page_records, cleaned_pages, chunk.chunk_index)
        chunk.output_part_path = export_part.local_docx_path
        chunk.progress_percent = 100.0
        chunk.status = ChunkStatus.completed
        chunk.current_stage = "completed"
        chunk.finished_at = datetime.utcnow()
        db.commit()
        finish_chunk_stage_run(db, part_run, status=ChunkStageStatus.completed, message="DOCX part generated")
        log_event(db, document_id, f"Chunk {chunk.chunk_index} completed", stage_key="part_generate", chunk_id=chunk.id)

    flush_log_events(db)
    pdf.close()

    update_job(db, job, stage=9, stage_key="appendix_extract", stage_name="Appendix Extraction", progress=92.0, message="Scanning parts for appendix/reference sections")
    appendix_path = extract_appendix_reference(document_id, db)

    update_job(db, job, stage=10, stage_key="merge_prep", stage_name="Merge Preparation", progress=93.0, message="Deduplicating content across chunks")
    _dedup_cleaned_chunk_texts(document_id, db)

    update_job(db, job, stage=10, stage_key="merge_prep", stage_name="Merge Preparation", progress=95.0, message="Writing manifest and indexing outputs")
    manifest_path = manifests_dir / "manifest.json"
    write_manifest(
        records=page_records,
        output_dir=manifests_dir,
        input_pdf=document.filename,
        start_page=start_page,
        end_page=end_page,
        pages_per_docx=config.pages_per_docx,
    )
    document.manifest_path = str(manifest_path)
    register_artifact(db, document_id, ArtifactType.manifest, "Manifest JSON", manifest_path)
    if appendix_path is not None:
        register_artifact(db, document_id, ArtifactType.appendix, "Appendix Reference", appendix_path)

    from app.services.manuscript_assembler import create_export_profile_defaults
    create_export_profile_defaults(db, document, config)

    document.status = DocumentStatus.merge_ready
    document.error_log = None
    job.status = JobStatus.completed
    job.completed_at = datetime.utcnow()
    flush_log_events(db)
    log_path = logs_dir / f"job_{job.id}.log"
    log_path.write_text("\n".join(event.message for event in db.query(ProjectEventLog).filter(ProjectEventLog.document_id == document_id).order_by(ProjectEventLog.created_at.asc()).all()), encoding="utf-8")
    register_artifact(db, document_id, ArtifactType.log, "Pipeline log", log_path)
    update_job(db, job, stage=11, stage_key="completed", stage_name="Completed", progress=100.0, message=f"Generated {len(created_parts)} DOCX parts", status=JobStatus.completed)
    db.commit()

    return {
        "document_id": document_id,
        "job_id": job_id,
        "page_count": total_pages,
        "parts_generated": len(created_parts),
        "manifest_path": str(manifest_path),
        "appendix_path": str(appendix_path) if appendix_path else None,
    }


def _annotate_blocks_with_spans(blocks: list, rich_pages: list) -> None:
    """
    Annotate transformed blocks with span metadata (bold/italic) from the source
    rich extraction. Matches each BODY block's original_text or text against the
    runs of LineItems in rich_pages and copies span info to block.spans.

    This enables `build_docx_from_blocks` to render per-run bold/italic emphasis
    even for AI-transformed blocks.
    """
    from pdf_to_clean_docx import LineItem
    span_index: list[dict] = []
    for _page_num, items in rich_pages:
        for item in items:
            if isinstance(item, LineItem) and item.runs:
                combined = "".join(r.text for r in item.runs)
                span_index.append({
                    "text": combined,
                    "spans": [{"text": r.text, "bold": r.bold, "italic": r.italic} for r in item.runs],
                })

    for block in blocks:
        block_type = getattr(block, "block_type", None) or (block.get("block_type") if isinstance(block, dict) else None)
        if block_type in ("H1", "H2", "H3", "TABLE"):
            continue
        src_text = ""
        if hasattr(block, "original_text"):
            src_text = (block.original_text or block.text or "").strip()
        elif isinstance(block, dict):
            src_text = (block.get("original_text") or block.get("text") or "").strip()
        if not src_text:
            continue
        best_spans: list = []
        for entry in span_index:
            entry_text = entry["text"].strip()
            if not entry_text:
                continue
            if entry_text in src_text:
                has_emphasis = any(s.get("bold") or s.get("italic") for s in entry["spans"])
                if has_emphasis and len(entry["spans"]) > len(best_spans):
                    best_spans = entry["spans"]
        if best_spans:
            if hasattr(block, "spans"):
                block.spans = best_spans
            elif isinstance(block, dict):
                block["spans"] = best_spans


_HEADING_BLOCK_TYPES = {"H1", "H2", "H3", "HEADING", "heading", "h1", "h2", "h3"}


def _dedup_blocks_before_render(blocks: list[dict]) -> list[dict]:
    """
    Deduplicate AI-transformed blocks BEFORE DOCX rendering.

    - BODY/TABLE blocks: near-duplicate removed using 4-gram Jaccard >= 60% threshold.
    - HEADING blocks: only removed on exact text match (case-insensitive) to avoid
      false-positive removal of intentionally repeated section titles.
    """
    from pdf_to_clean_docx import deduplicate_paragraphs

    def _block_type(b) -> str:
        return (b.get("block_type", "") if isinstance(b, dict) else getattr(b, "block_type", "")) or ""

    def _text(b) -> str:
        return (b.get("text", "") if isinstance(b, dict) else getattr(b, "text", "")) or ""

    body_texts = [
        _text(b) for b in blocks
        if _text(b).strip() and _block_type(b).upper() not in _HEADING_BLOCK_TYPES
    ]
    deduped_body = set(deduplicate_paragraphs([t for t in body_texts if t.strip()], similarity_threshold=0.60))

    seen_headings: set = set()
    result: list[dict] = []
    used_body_texts: set = set()
    for block in blocks:
        text = _text(block)
        stripped = text.strip()
        if not stripped:
            continue
        btype = _block_type(block).upper()
        if btype in _HEADING_BLOCK_TYPES:
            key = stripped.lower()
            if key in seen_headings:
                continue
            seen_headings.add(key)
            result.append(block)
        else:
            if stripped in deduped_body and stripped not in used_body_texts:
                used_body_texts.add(stripped)
                result.append(block)
            elif stripped not in deduped_body:
                pass
    return result


def _dedup_cleaned_chunk_texts(document_id: str, db: Session) -> None:
    """
    Merge-prep stage: cross-chunk paragraph deduplication over cleaned/transformed text files.

    For each chunk, reads its cleaned (or transformed) text, splits into paragraphs,
    removes near-duplicates relative to all previously seen paragraphs (Jaccard >= 60%),
    and writes the deduplicated text back to the same file.

    This runs at merge-prep stage so duplicates are eliminated in source text artifacts
    before any further post-processing.
    """
    from pdf_to_clean_docx import deduplicate_paragraphs
    chunks = db.query(Chunk).filter(Chunk.document_id == document_id).order_by(Chunk.chunk_index.asc()).all()
    global_paragraphs: list[str] = []

    for chunk in chunks:
        text_path_str = chunk.transformed_text_path or chunk.cleaned_text_path
        if not text_path_str:
            continue
        text_path = Path(text_path_str)
        if not text_path.exists():
            continue

        text = text_path.read_text(encoding="utf-8", errors="ignore")
        paragraphs = [p for p in text.split("\n\n") if p.strip()]

        deduped = deduplicate_paragraphs(global_paragraphs + paragraphs, similarity_threshold=0.60)
        new_paras = deduped[len(global_paragraphs):]
        global_paragraphs = deduped

        if len(new_paras) < len(paragraphs):
            log_event(db, document_id, f"Chunk {chunk.chunk_index}: dedup reduced {len(paragraphs)} -> {len(new_paras)} paragraphs", stage_key="merge_prep")
            text_path.write_text("\n\n".join(new_paras), encoding="utf-8")


def _assign_part_numbers(records: list[PageRecord], part_pages: list[tuple[int, str]], part_number: int) -> None:
    page_numbers = {page for page, _ in part_pages}
    for record in records:
        if record.source_page in page_numbers:
            record.part_number = part_number


def _flush_part(
    db: Session,
    document: Document,
    job: Job,
    outputs_dir: Path,
    part_pages: list[tuple[int, Any]],
    part_number: int,
    keep_page_markers: bool,
) -> ExportPart:
    output_path = outputs_dir / f"textbook_part_{part_number:02d}.docx"
    build_docx_part(
        part_number=part_number,
        part_pages=part_pages,
        output_path=output_path,
        keep_page_markers=keep_page_markers,
    )
    export_part = ExportPart(
        document_id=document.id,
        job_id=job.id,
        part_number=part_number,
        page_start=part_pages[0][0],
        page_end=part_pages[-1][0],
        page_count=len(part_pages),
        local_docx_path=str(output_path),
        filename=output_path.name,
        status=PartStatus.generated,
    )
    db.add(export_part)
    db.commit()
    db.refresh(export_part)
    register_artifact(db, document.id, ArtifactType.docx_part, f"DOCX part {part_number}", output_path)
    return export_part


def build_manifest_summary(document: Document) -> Optional[Dict[str, Any]]:
    if not document.manifest_path:
        return None
    manifest_path = Path(document.manifest_path)
    if not manifest_path.exists():
        return None
    data = json.loads(manifest_path.read_text(encoding="utf-8"))
    records = data.get("page_records", [])
    kept = [record for record in records if record.get("kept")]
    return {
        "input_pdf": data.get("input_pdf"),
        "start_page": data.get("start_page"),
        "end_page": data.get("end_page"),
        "pages_per_docx": data.get("pages_per_docx"),
        "total_records": len(records),
        "kept_pages": len(kept),
        "dropped_pages": len(records) - len(kept),
    }


def merge_document_parts(document: Document, parts: list[ExportPart]) -> Path:
    if not parts:
        raise ValueError("No parts available to merge")

    parts = sorted(parts, key=lambda item: item.part_number)
    master = DocxDocument(parts[0].local_docx_path)
    composer = Composer(master)
    for part in parts[1:]:
        composer.append(DocxDocument(part.local_docx_path))

    appendix_path = get_appendix_path(document.id)
    if appendix_path.exists():
        composer.append(DocxDocument(str(appendix_path)))

    merge_dir = get_merge_dir(document.id)
    merged_path = merge_dir / f"{Path(document.filename).stem}_merged.docx"
    composer.save(str(merged_path))

    _post_process_merged_docx(str(merged_path))

    return merged_path


def _build_front_matter_docx(profile, output_path: Path) -> None:
    """Create a front matter DOCX page from ExportProfile fields."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument()

    if profile.book_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(profile.book_title)
        run.bold = True
        run.font.size = Pt(24)

    if profile.subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(profile.subtitle)
        run.font.size = Pt(16)

    doc.add_paragraph()

    if profile.author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(profile.author)
        run.font.size = Pt(14)

    if profile.institution:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(profile.institution)
        run.font.size = Pt(12)

    doc.add_page_break()

    if profile.copyright_text:
        p = doc.add_paragraph()
        run = p.add_run(profile.copyright_text)
        run.font.size = Pt(10)
        doc.add_page_break()

    doc.save(str(output_path))


def merge_parts_ordered(document: Document, parts: list[ExportPart], profile=None, progress_callback=None) -> Path:
    """Merge ExportPart DOCX files in caller-supplied order, preserving formatting.
    If a profile with book_title or author is provided, a front matter page is prepended.
    After merge, a post-processing pass adds TOC + consistent header/footer to the merged DOCX.
    progress_callback(percent: float, message: str) is called periodically if provided."""
    if not parts:
        raise ValueError("No parts provided to merge")

    merge_dir = get_merge_dir(document.id)

    def _emit(pct: float, msg: str) -> None:
        if progress_callback:
            try:
                progress_callback(pct, msg)
            except Exception:
                pass

    _emit(5.0, "Preparing merge")

    has_front_matter = profile and any([
        getattr(profile, "book_title", None),
        getattr(profile, "subtitle", None),
        getattr(profile, "author", None),
        getattr(profile, "institution", None),
        getattr(profile, "copyright_text", None),
    ])

    total_parts = len(parts)

    if has_front_matter:
        _emit(10.0, "Building front matter")
        front_matter_path = merge_dir / f"{document.id}_front_matter.docx"
        _build_front_matter_docx(profile, front_matter_path)
        master = DocxDocument(str(front_matter_path))
        composer = Composer(master)
        for i, part in enumerate(parts):
            pct = 15.0 + ((i + 1) / total_parts) * 70.0
            _emit(pct, f"Stitching part {i + 1} of {total_parts}")
            composer.append(DocxDocument(part.local_docx_path))
    else:
        _emit(15.0, f"Loading part 1 of {total_parts}")
        master = DocxDocument(parts[0].local_docx_path)
        composer = Composer(master)
        for i, part in enumerate(parts[1:], start=1):
            pct = 15.0 + ((i + 1) / total_parts) * 70.0
            _emit(pct, f"Stitching part {i + 1} of {total_parts}")
            composer.append(DocxDocument(part.local_docx_path))

    _emit(85.0, "Saving merged DOCX")
    merged_path = merge_dir / f"{Path(document.filename).stem}_merged.docx"
    composer.save(str(merged_path))

    _emit(90.0, "Applying publication-quality post-processing (TOC + headers/footers)")
    _post_process_merged_docx(str(merged_path), profile)

    _emit(100.0, "Merge complete")
    return merged_path


def _post_process_merged_docx(docx_path: str, profile=None) -> None:
    """
    Post-processing pass on the merged DOCX:
    1. Set page layout (A4, 2.54cm margins) on all sections.
    2. Add/replace header (book title, right-aligned italic) and footer (centred page number).
    3. Insert a Table of Contents page before body content if profile.include_toc is True.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _BODY_FONT = "Times New Roman"

    doc = DocxDocument(docx_path)

    book_title = ""
    include_toc = True
    if profile:
        book_title = (getattr(profile, "book_title", "") or "").strip()
        include_toc = bool(getattr(profile, "include_toc", True))

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.header_distance = Cm(1.27)
        section.footer_distance = Cm(1.27)
        section.different_first_page_header_footer = False

        header = section.header
        for existing_para in list(header.paragraphs):
            existing_para._element.getparent().remove(existing_para._element)
        hp = header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if book_title:
            run = hp.add_run(book_title)
            run.font.name = _BODY_FONT
            run.font.size = Pt(9)
            run.italic = True
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        footer = section.footer
        for existing_para in list(footer.paragraphs):
            existing_para._element.getparent().remove(existing_para._element)
        fp = footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        run.font.name = _BODY_FONT
        run.font.size = Pt(9)
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        run._r.append(fld)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run._r.append(instr)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)

    from pdf_to_clean_docx import deduplicate_paragraphs
    import logging as _log_mod
    _pp_log = _log_mod.getLogger(__name__)
    body_para_elements = [
        p for p in doc.paragraphs
        if p.text.strip()
        and not p.style.name.startswith("Heading")
        and p.style.name not in ("Title", "Subtitle", "TOC 1", "TOC 2", "TOC 3")
        and len(p.text.split()) >= 4
    ]
    body_texts = [p.text for p in body_para_elements]
    deduped_texts = deduplicate_paragraphs(body_texts, similarity_threshold=0.60)
    deduped_indices: set = set()
    for kept_text in deduped_texts:
        for idx, t in enumerate(body_texts):
            if t == kept_text and idx not in deduped_indices:
                deduped_indices.add(idx)
                break
    removed_count = 0
    for idx, p in enumerate(body_para_elements):
        if idx not in deduped_indices:
            p._element.getparent().remove(p._element)
            removed_count += 1
    if removed_count:
        _pp_log.info(f"Post-merge dedup: removed {removed_count} duplicate body paragraphs from merged DOCX")

    if include_toc:
        toc_heading = doc.add_paragraph()
        toc_heading._p.getparent().remove(toc_heading._p)

        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        body = doc.element.body
        first_para = body[0] if len(body) else None

        toc_h_p = OxmlElement("w:p")
        toc_h_r = OxmlElement("w:r")
        toc_h_rpr = OxmlElement("w:rPr")
        toc_h_b = OxmlElement("w:b")
        toc_h_sz = OxmlElement("w:sz")
        toc_h_sz.set(qn("w:val"), "36")
        toc_h_rpr.append(toc_h_b)
        toc_h_rpr.append(toc_h_sz)
        toc_h_r.append(toc_h_rpr)
        toc_h_t = OxmlElement("w:t")
        toc_h_t.text = "Table of Contents"
        toc_h_r.append(toc_h_t)
        toc_h_p.append(toc_h_r)

        toc_p = OxmlElement("w:p")
        toc_r = OxmlElement("w:r")
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        toc_r.append(fld_begin)
        instr_r = OxmlElement("w:r")
        instr_el = OxmlElement("w:instrText")
        instr_el.set(qn("xml:space"), "preserve")
        instr_el.text = ' TOC \\o "1-3" \\h \\z \\u '
        instr_r.append(instr_el)
        fld_end2 = OxmlElement("w:fldChar")
        fld_end2.set(qn("w:fldCharType"), "end")
        fld_end_r = OxmlElement("w:r")
        fld_end_r.append(fld_end2)
        toc_p.append(toc_r)
        toc_p.append(instr_r)
        toc_p.append(fld_end_r)

        page_break_p = OxmlElement("w:p")
        pb_r = OxmlElement("w:r")
        pb_br = OxmlElement("w:br")
        pb_br.set(qn("w:type"), "page")
        pb_r.append(pb_br)
        page_break_p.append(pb_r)

        if first_para is not None:
            body.insert(0, page_break_p)
            body.insert(0, toc_p)
            body.insert(0, toc_h_p)
        else:
            body.append(toc_h_p)
            body.append(toc_p)
            body.append(page_break_p)

    doc.save(docx_path)


def validate_merge(document: Document, parts: list[ExportPart], chunks: list[Chunk]) -> Dict[str, Any]:
    missing_chunks = [chunk.chunk_index for chunk in chunks if chunk.status != ChunkStatus.completed]
    missing_parts = [chunk.chunk_index for chunk in chunks if chunk.status == ChunkStatus.completed and not chunk.output_part_path]
    return {
        "ok": not missing_chunks and not missing_parts and bool(parts),
        "missing_chunks": missing_chunks,
        "missing_parts": missing_parts,
        "part_count": len(parts),
        "chunk_count": len(chunks),
        "status": document.status.value if hasattr(document.status, "value") else str(document.status),
    }


def build_system_health(db: Session) -> Dict[str, Any]:
    documents = db.query(Document).all()
    chunks = db.query(Chunk).all()
    failed = [doc for doc in documents if getattr(doc.status, "value", doc.status) == "FAILED"]
    storage_root = Path(settings.PROJECTS_DIR)
    total_bytes = 0
    if storage_root.exists():
        for root, _, files in os.walk(storage_root):
            for name in files:
                try:
                    total_bytes += (Path(root) / name).stat().st_size
                except OSError:
                    pass
    return {
        "project_count": len(documents),
        "processing_count": sum(1 for doc in documents if getattr(doc.status, "value", doc.status) == "PROCESSING"),
        "failed_count": len(failed),
        "ready_count": sum(1 for doc in documents if getattr(doc.status, "value", doc.status) in {"READY", "MERGE_READY", "MERGED"}),
        "chunk_count": len(chunks),
        "failed_chunks": sum(1 for chunk in chunks if getattr(chunk.status, "value", chunk.status) == "FAILED"),
        "storage_bytes": total_bytes,
    }


def rerun_chunk(document_id: str, chunk_id: str, db: Session, stage_key: str = "extract") -> Dict[str, Any]:
    document = db.query(Document).filter(Document.id == document_id).first()
    chunk = db.query(Chunk).filter(Chunk.id == chunk_id, Chunk.document_id == document_id).first()
    config = _get_cached_document_config(document_id, db)
    if not document or not chunk or not config:
        raise ValueError("Document, chunk, or config not found")

    outputs_dir = get_outputs_dir(document_id)
    chunks_dir = get_chunks_dir(document_id)
    pdf = fitz.open(document.local_storage_path)
    try:
        chunk.retry_count = (chunk.retry_count or 0) + 1
        chunk.status = ChunkStatus.processing
        chunk.current_stage = stage_key
        chunk.progress_percent = 5.0
        chunk.error_log = None
        chunk.started_at = datetime.utcnow()
        db.commit()
        log_event(db, document_id, f"Rerunning chunk {chunk.chunk_index} from stage {stage_key}", stage_key=stage_key, chunk_id=chunk.id)

        cleaned_pages = []
        if stage_key == "part_generate" and chunk.cleaned_text_path and Path(chunk.cleaned_text_path).exists():
            blocks = Path(chunk.cleaned_text_path).read_text(encoding="utf-8", errors="ignore").split("===== PAGE ")
            for block in blocks:
                block = block.strip()
                if not block:
                    continue
                newline_idx = block.find("\n")
                page_num = int(block[:newline_idx].strip())
                cleaned_pages.append((page_num, block[newline_idx + 1 :].strip()))
        else:
            rerun_extract = create_chunk_stage_run(db, chunk, "extract", "Extract Pages", message="Rerun extraction")
            raw_lines = []
            rerun_table_headers: set = set()
            for source_page in range(chunk.page_start, chunk.page_end + 1):
                page_obj = pdf[source_page - 1]
                try:
                    items = extract_page_rich(page_obj)
                    rich_plain = content_items_to_plain_text(items)
                    raw_text = rich_plain if rich_plain.strip() else page_obj.get_text("text")
                except Exception:
                    raw_text = page_obj.get_text("text")
                raw_lines.append(f"===== PAGE {source_page} =====\n{raw_text}")
                cleaned = clean_page_text(raw_text, seen_table_headers=rerun_table_headers)
                if cleaned.strip():
                    cleaned_pages.append((source_page, cleaned))

            raw_path = chunks_dir / f"chunk_{chunk.chunk_index:04d}_raw.txt"
            raw_path.write_text("\n\n".join(raw_lines), encoding="utf-8")
            chunk.raw_text_path = str(raw_path)
            finish_chunk_stage_run(db, rerun_extract, status=ChunkStageStatus.completed, message="Rerun extraction complete")
            chunk.current_stage = "clean_pass_1"
            chunk.progress_percent = 45.0
            db.commit()
            rerun_clean = create_chunk_stage_run(db, chunk, "clean_pass_1", "Cleanup Pass 1", message="Rerun cleanup")

            cleaned_path = chunks_dir / f"chunk_{chunk.chunk_index:04d}_cleaned.txt"
            cleaned_path.write_text("\n\n".join(f"===== PAGE {p} =====\n{text}" for p, text in cleaned_pages), encoding="utf-8")
            chunk.cleaned_text_path = str(cleaned_path)
            finish_chunk_stage_run(db, rerun_clean, status=ChunkStageStatus.completed, message="Rerun cleanup complete")

        chunk.current_stage = "part_generate"
        chunk.progress_percent = 70.0
        db.commit()
        rerun_part = create_chunk_stage_run(db, chunk, "part_generate", "Generate DOCX Part", message="Rerun DOCX generation")

        existing_part = db.query(ExportPart).filter(ExportPart.document_id == document_id, ExportPart.part_number == chunk.chunk_index).first()
        output_path = outputs_dir / f"textbook_part_{chunk.chunk_index:02d}.docx"
        build_docx_part(chunk.chunk_index, cleaned_pages, output_path, config.keep_page_markers)
        if existing_part:
            existing_part.page_start = cleaned_pages[0][0] if cleaned_pages else chunk.page_start
            existing_part.page_end = cleaned_pages[-1][0] if cleaned_pages else chunk.page_end
            existing_part.page_count = len(cleaned_pages)
            existing_part.local_docx_path = str(output_path)
            existing_part.filename = output_path.name
            existing_part.status = PartStatus.generated
        else:
            db.add(
                ExportPart(
                    document_id=document.id,
                    job_id=document.jobs[0].id if document.jobs else str(uuid.uuid4()),
                    part_number=chunk.chunk_index,
                    page_start=cleaned_pages[0][0] if cleaned_pages else chunk.page_start,
                    page_end=cleaned_pages[-1][0] if cleaned_pages else chunk.page_end,
                    page_count=len(cleaned_pages),
                    local_docx_path=str(output_path),
                    filename=output_path.name,
                    status=PartStatus.generated,
                )
            )
        chunk.output_part_path = str(output_path)
        chunk.status = ChunkStatus.completed
        chunk.current_stage = "completed"
        chunk.progress_percent = 100.0
        chunk.finished_at = datetime.utcnow()
        db.commit()
        finish_chunk_stage_run(db, rerun_part, status=ChunkStageStatus.completed, message="Rerun DOCX generation complete")
        log_event(db, document_id, f"Chunk {chunk.chunk_index} rerun completed", stage_key="part_generate", chunk_id=chunk.id, flush=True)
        return {"status": "success", "chunk_id": chunk.id}
    finally:
        pdf.close()


def extract_appendix_reference(document_id: str, db: Session) -> Optional[Path]:
    headings = ("GLOSSARY", "HOUSEHOLD MEASURES", "CONVERSION CHART", "KCAL EQUIVALENCE", "REFERENCE INDEX")
    parts = db.query(ExportPart).filter(ExportPart.document_id == document_id).order_by(ExportPart.part_number.asc()).all()
    if not parts:
        return None

    appendix_doc = DocxDocument()
    in_appendix = False
    found_any = False
    for part in parts:
        if not Path(part.local_docx_path).exists():
            continue
        doc = DocxDocument(part.local_docx_path)
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if any(text.upper().startswith(heading) for heading in headings):
                in_appendix = True
            if in_appendix:
                appendix_doc.add_paragraph(text)
                found_any = True

    if not found_any:
        return None

    appendix_path = get_appendix_path(document_id)
    appendix_doc.save(str(appendix_path))
    return appendix_path


def initialize_pipeline_fanout(document_id: str, job_id: str, db: Session) -> list[str]:
    document = db.query(Document).filter(Document.id == document_id).first()
    job = db.query(Job).filter(Job.id == job_id).first()
    config = _get_cached_document_config(document_id, db)
    if not document or not job or not config:
        raise ValueError("Document, job, or config not found")

    project_dir = get_project_dir(document_id)
    get_outputs_dir(document_id)
    get_chunks_dir(document_id)
    get_manifests_dir(document_id)
    get_logs_dir(document_id)
    get_merge_dir(document_id)

    document.storage_root = str(project_dir)
    document.status = DocumentStatus.processing
    job.status = JobStatus.in_progress
    db.commit()

    update_job(db, job, stage=2, stage_key="inspect", stage_name="Inspecting PDF", progress=5.0, message="Reading PDF metadata")
    meta = inspect_pdf(document.local_storage_path)
    document.page_count = int(meta["page_count"])
    document.status = DocumentStatus.inspected
    db.commit()
    log_event(db, document_id, f"PDF inspected: {document.page_count} pages", stage_key="inspect")

    start_page = max(1, int(config.start_page))
    end_page = min(int(document.page_count or 0), int(config.end_page))
    update_job(db, job, stage=3, stage_key="plan", stage_name="Planning Chunks", progress=10.0, message=f"Planning chunks for pages {start_page}-{end_page}")

    db.query(ExportPart).filter(ExportPart.document_id == document_id).delete(synchronize_session=False)
    db.query(Chunk).filter(Chunk.document_id == document_id).delete(synchronize_session=False)
    db.query(Artifact).filter(Artifact.document_id == document_id).delete(synchronize_session=False)
    db.commit()

    chunk_ids: list[str] = []
    page_cursor = start_page
    chunk_index = 1
    while page_cursor <= end_page:
        chunk_end = min(end_page, page_cursor + int(config.pages_per_docx) - 1)
        chunk = Chunk(
            document_id=document_id,
            chunk_index=chunk_index,
            page_start=page_cursor,
            page_end=chunk_end,
            page_count=chunk_end - page_cursor + 1,
            status=ChunkStatus.queued,
            current_stage="planned",
        )
        db.add(chunk)
        db.commit()
        db.refresh(chunk)
        chunk_ids.append(chunk.id)
        chunk_index += 1
        page_cursor = chunk_end + 1

    document.status = DocumentStatus.planned
    db.commit()
    log_event(db, document_id, f"Planned {len(chunk_ids)} chunks", stage_key="plan", flush=True)
    return chunk_ids


def process_single_chunk(document_id: str, chunk_id: str, db: Session) -> Dict[str, Any]:
    document = db.query(Document).filter(Document.id == document_id).first()
    chunk = db.query(Chunk).filter(Chunk.id == chunk_id, Chunk.document_id == document_id).first()
    config = _get_cached_document_config(document_id, db)
    job = db.query(Job).filter(Job.document_id == document_id).order_by(Job.created_at.desc()).first()
    if not document or not chunk or not config or not job:
        raise ValueError("Document, chunk, config, or job not found")

    outputs_dir = get_outputs_dir(document_id)
    chunks_dir = get_chunks_dir(document_id)
    total_chunks = max(1, db.query(Chunk).filter(Chunk.document_id == document_id).count())

    pdf = fitz.open(document.local_storage_path)
    try:
        chunk.status = ChunkStatus.processing
        chunk.current_stage = "extract"
        chunk.started_at = datetime.utcnow()
        db.commit()
        log_event(db, document_id, f"Processing chunk {chunk.chunk_index} ({chunk.page_start}-{chunk.page_end})", stage_key="extract", chunk_id=chunk.id)
        extract_run = create_chunk_stage_run(db, chunk, "extract", "Extract Pages", message="Reading PDF pages")

        raw_lines = []
        cleaned_pages = []
        rich_pages: list = []
        single_chunk_table_headers: set = set()
        for source_page in range(int(chunk.page_start), int(chunk.page_end) + 1):
            page_obj = pdf[source_page - 1]
            try:
                items = extract_page_rich(page_obj)
                rich_plain = content_items_to_plain_text(items)
                raw_text = rich_plain if rich_plain.strip() else page_obj.get_text("text")
            except Exception:
                items = []
                raw_text = page_obj.get_text("text")
            raw_lines.append(f"===== PAGE {source_page} =====\n{raw_text}")
            cleaned = clean_page_text(raw_text, seen_table_headers=single_chunk_table_headers)
            if cleaned.strip():
                cleaned_pages.append((source_page, cleaned))
                rich_pages.append((source_page, items))

        raw_path = chunks_dir / f"chunk_{int(chunk.chunk_index):04d}_raw.txt"
        raw_path.write_text("\n\n".join(raw_lines), encoding="utf-8")
        chunk.raw_text_path = str(raw_path)
        register_artifact(db, document_id, ArtifactType.chunk_text, f"Chunk {chunk.chunk_index} raw text", raw_path, chunk.id)
        finish_chunk_stage_run(db, extract_run, status=ChunkStageStatus.completed, message="Raw extraction complete")

        update_job(db, job, stage=5, stage_key="clean_pass_1", stage_name="Cleanup Pass 1", progress=15.0 + ((int(chunk.chunk_index) - 1) / total_chunks) * 60.0, message=f"Cleanup pass 1 for chunk {chunk.chunk_index}/{total_chunks}")
        chunk.current_stage = "clean_pass_1"
        db.commit()
        clean1_run = create_chunk_stage_run(db, chunk, "clean_pass_1", "Cleanup Pass 1", message="Applying deterministic cleanup rules")

        cleaned_text = "\n\n".join(f"===== PAGE {page} =====\n{text}" for page, text in cleaned_pages)
        cleaned_path = chunks_dir / f"chunk_{int(chunk.chunk_index):04d}_cleaned.txt"
        cleaned_path.write_text(cleaned_text, encoding="utf-8")
        chunk.cleaned_text_path = str(cleaned_path)
        register_artifact(db, document_id, ArtifactType.cleaned_text, f"Chunk {chunk.chunk_index} cleaned text", cleaned_path, chunk.id)
        finish_chunk_stage_run(db, clean1_run, status=ChunkStageStatus.completed, message="Chunk text cleaned")

        chunk.current_stage = "clean_pass_2"
        db.commit()
        clean2_run = create_chunk_stage_run(db, chunk, "clean_pass_2", "Cleanup Pass 2", message="Second cleanup checkpoint")
        finish_chunk_stage_run(db, clean2_run, status=ChunkStageStatus.completed, message="Second cleanup pass complete")

        chunk.current_stage = "final_normalize"
        db.commit()
        normalize_run = create_chunk_stage_run(db, chunk, "final_normalize", "Final Normalize", message="Normalizing chunk output")
        finish_chunk_stage_run(db, normalize_run, status=ChunkStageStatus.completed, message="Normalization complete")

        update_job(db, job, stage=8, stage_key="ai_transform", stage_name="AI Transformation", progress=20.0 + (int(chunk.chunk_index) / total_chunks) * 40.0, message=f"AI transformation for chunk {chunk.chunk_index}/{total_chunks}")
        chunk.current_stage = "ai_transform"
        db.commit()
        transform_run = create_chunk_stage_run(db, chunk, "ai_transform", "AI Transformation", message="Running two-pass AI content transformation")

        transformed_blocks = None
        try:
            from app.core.ai_pool import get_ai_pool
            from app.services.ai_transformer import (
                build_transformed_text,
                transform_chunk_text,
            )
            import dataclasses as _dc
            import json as _json

            ai_pool = get_ai_pool()
            if ai_pool.available():
                def _transform_progress(stage_label: str, pct: float) -> None:
                    try:
                        prog = 20.0 + (int(chunk.chunk_index) / total_chunks) * 40.0 + pct * 0.40
                        update_job(db, job, stage=8, stage_key="ai_transform", stage_name="AI Transformation",
                                   progress=min(prog, 60.0), message=f"[{stage_label}] chunk {chunk.chunk_index}")
                    except Exception:
                        pass

                log_event(db, document_id, f"AI transform starting for chunk {chunk.chunk_index}", stage_key="ai_transform", chunk_id=chunk.id)
                blocks, stats = transform_chunk_text(cleaned_text, ai_pool, progress_cb=_transform_progress)

                if blocks:
                    transformed_txt = build_transformed_text(blocks)
                    transformed_path = chunks_dir / f"chunk_{int(chunk.chunk_index):04d}_transformed.txt"
                    transformed_path.write_text(transformed_txt, encoding="utf-8")
                    chunk.transformed_text_path = str(transformed_path)
                    chunk.transformed_text = transformed_txt
                    chunk.transform_stats = _dc.asdict(stats)

                    transformed_json_path = chunks_dir / f"chunk_{int(chunk.chunk_index):04d}_transformed.json"
                    transformed_json_path.write_text(
                        _json.dumps([_dc.asdict(b) for b in blocks], ensure_ascii=False, indent=2),
                        encoding="utf-8",
                    )

                    register_artifact(db, document_id, ArtifactType.transformed_text, f"Chunk {chunk.chunk_index} transformed text", transformed_path, chunk.id)
                    register_artifact(db, document_id, ArtifactType.transformed_json, f"Chunk {chunk.chunk_index} transformed JSON", transformed_json_path, chunk.id)

                    _annotate_blocks_with_spans(blocks, rich_pages)
                    transformed_blocks = blocks
                    log_event(db, document_id, f"AI transform done: {stats.total_blocks} blocks, {stats.rewritten_blocks} rewritten, {stats.fallback_blocks} fallbacks", stage_key="ai_transform", chunk_id=chunk.id)
                else:
                    log_event(db, document_id, "AI transform produced no blocks — using cleaned text for DOCX", stage_key="ai_transform", chunk_id=chunk.id)
            else:
                log_event(db, document_id, "No AI keys configured — skipping ai_transform stage", stage_key="ai_transform", chunk_id=chunk.id)
        except Exception as ai_exc:
            log_event(db, document_id, f"AI transform error (non-fatal): {ai_exc}", level="WARNING", stage_key="ai_transform", chunk_id=chunk.id)

        finish_chunk_stage_run(db, transform_run, status=ChunkStageStatus.completed, message="AI transformation complete")

        update_job(db, job, stage=9, stage_key="part_generate", stage_name="Generating DOCX Parts", progress=25.0 + (int(chunk.chunk_index) / total_chunks) * 60.0, message=f"Generating DOCX part for chunk {chunk.chunk_index}/{total_chunks}")
        chunk.current_stage = "part_generate"
        db.commit()
        part_run = create_chunk_stage_run(db, chunk, "part_generate", "Generate DOCX Part", message="Building DOCX for chunk")

        export_part = db.query(ExportPart).filter(ExportPart.document_id == document_id, ExportPart.part_number == int(chunk.chunk_index)).first()
        output_path = outputs_dir / f"textbook_part_{int(chunk.chunk_index):02d}.docx"
        if transformed_blocks:
            deduped_blocks = _dedup_blocks_before_render(transformed_blocks)
            log_event(db, document_id, f"Dedup: {len(transformed_blocks)} -> {len(deduped_blocks)} blocks for chunk {chunk.chunk_index}", stage_key="part_generate", chunk_id=chunk.id)
            build_docx_from_blocks(int(chunk.chunk_index), deduped_blocks, output_path)
            log_event(db, document_id, f"DOCX part {chunk.chunk_index} built from {len(deduped_blocks)} AI-transformed blocks", stage_key="part_generate", chunk_id=chunk.id)
        else:
            build_docx_part(int(chunk.chunk_index), cleaned_pages, output_path, bool(config.keep_page_markers))
        if export_part:
            export_part.page_start = cleaned_pages[0][0] if cleaned_pages else int(chunk.page_start)
            export_part.page_end = cleaned_pages[-1][0] if cleaned_pages else int(chunk.page_end)
            export_part.page_count = len(cleaned_pages)
            export_part.local_docx_path = str(output_path)
            export_part.filename = output_path.name
            export_part.status = PartStatus.generated
        else:
            db.add(ExportPart(document_id=document_id, job_id=job.id, part_number=int(chunk.chunk_index), page_start=cleaned_pages[0][0] if cleaned_pages else int(chunk.page_start), page_end=cleaned_pages[-1][0] if cleaned_pages else int(chunk.page_end), page_count=len(cleaned_pages), local_docx_path=str(output_path), filename=output_path.name, status=PartStatus.generated))
        db.commit()
        register_artifact(db, document_id, ArtifactType.docx_part, f"DOCX part {chunk.chunk_index}", output_path)

        chunk.output_part_path = str(output_path)
        chunk.progress_percent = 100.0
        chunk.status = ChunkStatus.completed
        chunk.current_stage = "completed"
        chunk.finished_at = datetime.utcnow()
        db.commit()
        finish_chunk_stage_run(db, part_run, status=ChunkStageStatus.completed, message="DOCX part generated")
        log_event(db, document_id, f"Chunk {chunk.chunk_index} completed", stage_key="part_generate", chunk_id=chunk.id, flush=True)
        return {"status": "success", "chunk_id": chunk.id}
    except Exception as exc:
        chunk.status = ChunkStatus.failed
        chunk.error_log = str(exc)
        chunk.finished_at = datetime.utcnow()
        db.commit()
        log_event(db, document_id, f"Chunk {chunk.chunk_index} failed: {exc}", level="ERROR", stage_key=chunk.current_stage, chunk_id=chunk.id, flush=True)
        raise
    finally:
        pdf.close()


def finalize_pipeline(document_id: str, job_id: str, db: Session) -> Dict[str, Any]:
    document = db.query(Document).filter(Document.id == document_id).first()
    job = db.query(Job).filter(Job.id == job_id).first()
    config = _get_cached_document_config(document_id, db)
    if not document or not job or not config:
        raise ValueError("Document, job, or config not found")

    if bool(document.finalize_in_progress):
        return {"status": "skipped", "message": "Finalize already running"}

    chunks = db.query(Chunk).filter(Chunk.document_id == document_id).order_by(Chunk.chunk_index.asc()).all()
    if not chunks or any(chunk.status != ChunkStatus.completed for chunk in chunks):
        return {"status": "waiting", "message": "Chunks still pending"}

    if getattr(job.status, "value", job.status) == "COMPLETED" and document.manifest_path:
        return {"status": "already_completed", "message": "Pipeline already finalized"}

    document.finalize_in_progress = True
    db.commit()

    manifests_dir = get_manifests_dir(document_id)
    logs_dir = get_logs_dir(document_id)
    try:
        update_job(db, job, stage=10, stage_key="appendix_extract", stage_name="Appendix Extraction", progress=92.0, message="Scanning parts for appendix/reference sections")
        appendix_path = extract_appendix_reference(document_id, db)

        update_job(db, job, stage=11, stage_key="merge_prep", stage_name="Merge Preparation", progress=93.0, message="Deduplicating content across chunks")
        _dedup_cleaned_chunk_texts(document_id, db)

        update_job(db, job, stage=11, stage_key="merge_prep", stage_name="Merge Preparation", progress=95.0, message="Writing manifest and indexing outputs")

        pdf = fitz.open(document.local_storage_path)
        try:
            start_page = max(1, int(config.start_page))
            end_page = min(len(pdf), int(config.end_page))
            records: list[PageRecord] = []
            for page_number in range(start_page, end_page + 1):
                cleaned = clean_page_text(pdf[page_number - 1].get_text("text"))
                part_number = None
                for chunk in chunks:
                    if int(chunk.page_start) <= page_number <= int(chunk.page_end):
                        part_number = int(chunk.chunk_index)
                        break
                record = PageRecord(source_page=page_number, cleaned_char_count=len(cleaned), kept=bool(cleaned.strip()), part_number=part_number)
                records.append(record)
        finally:
            pdf.close()

        manifest_path = manifests_dir / "manifest.json"
        write_manifest(records=records, output_dir=manifests_dir, input_pdf=document.filename, start_page=max(1, int(config.start_page)), end_page=min(int(document.page_count or 0), int(config.end_page)), pages_per_docx=int(config.pages_per_docx))
        document.manifest_path = str(manifest_path)
        register_artifact(db, document_id, ArtifactType.manifest, "Manifest JSON", manifest_path)
        if appendix_path is not None:
            register_artifact(db, document_id, ArtifactType.appendix, "Appendix Reference", appendix_path)

        log_path = logs_dir / f"job_{job.id}.log"
        log_path.write_text("\n".join(event.message for event in db.query(ProjectEventLog).filter(ProjectEventLog.document_id == document_id).order_by(ProjectEventLog.created_at.asc()).all()), encoding="utf-8")
        register_artifact(db, document_id, ArtifactType.log, "Pipeline log", log_path)

        document.status = DocumentStatus.merge_ready
        document.error_log = None
        job.status = JobStatus.completed
        job.completed_at = datetime.utcnow()
        update_job(db, job, stage=12, stage_key="completed", stage_name="Completed", progress=100.0, message=f"Generated {len(chunks)} DOCX parts", status=JobStatus.completed)
        log_event(db, document_id, "Pipeline finalized and ready for merge", stage_key="completed", flush=True)
        db.commit()
        return {"status": "success", "parts_generated": len(chunks), "manifest_path": str(manifest_path)}
    finally:
        document.finalize_in_progress = False
        db.commit()
        _evict_document_config(document_id)
