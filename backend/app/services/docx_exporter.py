"""
docx_exporter.py — Professional DOCX export meeting Indian university publishing standards.

Features:
- Named Word styles (Heading 1-3, Body Text, Caption, Table Grid, etc.)
- Times New Roman 12pt body, Arial headings, 1.5 line spacing
- A4 page with 2.54cm margins, justified text, first-line paragraph indent
- Header (book title, top right) and footer (page number, bottom center)
- Word built-in TOC field (auto-populated when document is opened in Word)
- Professional front matter: title page, copyright, dedication, preface
- Back matter: References and Index placeholders
- Chapter headings on new pages with "Chapter N: Title" formatting
- Tables with Table Grid style, bold shaded header rows, auto-fit columns
"""

import os
import re
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.styles.style import ParagraphStyle


# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------
_HEADER_BG = RGBColor(0x2E, 0x2E, 0x2E)   # dark charcoal for table headers
_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)   # white text for table headers
_RULE_COLOR = RGBColor(0x80, 0x80, 0x80)  # thin rule colour


# ---------------------------------------------------------------------------
# Style definitions
# ---------------------------------------------------------------------------
_BODY_FONT = "Times New Roman"
_HEAD_FONT = "Arial"


def _ensure_style(doc: DocxDocument, name: str, base_name: Optional[str] = None) -> ParagraphStyle:
    """Return existing named style or create it."""
    if name in doc.styles:
        return doc.styles[name]
    base = doc.styles[base_name] if base_name and base_name in doc.styles else doc.styles["Normal"]
    style = doc.styles.add_style(name, 1)  # 1 = paragraph style
    style.base_style = base
    return style


def _define_styles(doc: DocxDocument) -> None:
    """Apply professional style definitions to the document."""

    # --- Normal / Body Text --------------------------------------------------
    normal = doc.styles["Normal"]
    normal.font.name = _BODY_FONT
    normal.font.size = Pt(12)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Body Text style (alias with first-line indent & justified alignment)
    bt = _ensure_style(doc, "Body Text", "Normal")
    bt.font.name = _BODY_FONT
    bt.font.size = Pt(12)
    bt.paragraph_format.first_line_indent = Cm(1.27)
    bt.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bt.paragraph_format.space_before = Pt(0)
    bt.paragraph_format.space_after = Pt(6)
    bt.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # --- Headings ------------------------------------------------------------
    for level, size, space_before, space_after, bold, italic in [
        (1, 18, 24, 12, True,  False),
        (2, 14, 18,  8, True,  False),
        (3, 12, 12,  6, True,  True),
    ]:
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            h = doc.styles[style_name]
        else:
            h = doc.styles.add_style(style_name, 1)
            h.base_style = doc.styles["Normal"]
        h.font.name = _HEAD_FONT
        h.font.size = Pt(size)
        h.font.bold = bold
        h.font.italic = italic
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        h.paragraph_format.space_before = Pt(space_before)
        h.paragraph_format.space_after = Pt(space_after)
        h.paragraph_format.keep_with_next = True
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        h.paragraph_format.line_spacing = Pt(size * 1.4)

    # Heading 1: page-break-before is set per-paragraph during rendering

    # --- Caption -------------------------------------------------------------
    cap = _ensure_style(doc, "Caption", "Normal")
    cap.font.name = _BODY_FONT
    cap.font.size = Pt(10)
    cap.font.italic = True
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(8)

    # --- Front Matter heading -------------------------------------------------
    fm = _ensure_style(doc, "Front Matter Heading", "Normal")
    fm.font.name = _HEAD_FONT
    fm.font.size = Pt(16)
    fm.font.bold = True
    fm.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fm.paragraph_format.space_before = Pt(12)
    fm.paragraph_format.space_after = Pt(12)

    # --- Chapter Opener -------------------------------------------------------
    co = _ensure_style(doc, "Chapter Opener", "Normal")
    co.font.name = _HEAD_FONT
    co.font.size = Pt(20)
    co.font.bold = True
    co.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    co.paragraph_format.space_before = Pt(36)
    co.paragraph_format.space_after = Pt(24)
    co.paragraph_format.keep_with_next = True


# ---------------------------------------------------------------------------
# Page layout
# ---------------------------------------------------------------------------

def _set_page_layout(doc: DocxDocument, profile) -> None:
    section = doc.sections[0]
    page_size = getattr(profile, "page_size", "A4") or "A4"
    if page_size == "A4":
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    elif page_size == "US Letter":
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.94)
    # else: leave as default

    top    = getattr(profile, "margin_top_cm",    2.54) or 2.54
    bottom = getattr(profile, "margin_bottom_cm", 2.54) or 2.54
    left   = getattr(profile, "margin_left_cm",   2.54) or 2.54
    right  = getattr(profile, "margin_right_cm",  2.54) or 2.54

    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)
    section.different_first_page_header_footer = True


def _add_header_footer(doc: DocxDocument, profile) -> None:
    """Add book title (top-right header) and page number (bottom-center footer)."""
    book_title = (getattr(profile, "book_title", "") or "").strip()

    section = doc.sections[0]

    # -- body header (not first page) --
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.clear()
    if book_title:
        run = hp.add_run(book_title)
        run.font.name = _BODY_FONT
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # -- body footer (not first page) --
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()
    # Word PAGE field
    run = fp.add_run()
    run.font.name = _BODY_FONT
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    # -- first-page header/footer (blank for title page) --
    first_header = section.first_page_header
    if first_header.paragraphs:
        first_header.paragraphs[0].clear()
    first_footer = section.first_page_footer
    if first_footer.paragraphs:
        first_footer.paragraphs[0].clear()


# ---------------------------------------------------------------------------
# TOC field
# ---------------------------------------------------------------------------

def _add_toc(doc: DocxDocument, profile) -> None:
    toc_levels = getattr(profile, "toc_heading_levels", 2) or 2
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = rf' TOC \o "1-{toc_levels}" \h \z \u '
    run._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


# ---------------------------------------------------------------------------
# Page break helpers
# ---------------------------------------------------------------------------

def _page_break_before_para(para) -> None:
    """Set page-break-before on the paragraph via pPr element."""
    pPr = para._p.get_or_add_pPr()
    pb = OxmlElement("w:pageBreakBefore")
    pb.set(qn("w:val"), "true")
    pPr.append(pb)


def _add_hard_page_break(doc: DocxDocument) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ---------------------------------------------------------------------------
# Front matter
# ---------------------------------------------------------------------------

def _add_title_page(doc: DocxDocument, profile) -> None:
    title     = (getattr(profile, "book_title",  "") or "Untitled Manuscript").strip()
    subtitle  = (getattr(profile, "subtitle",    "") or "").strip()
    author    = (getattr(profile, "author",      "") or "").strip()
    edition   = (getattr(profile, "edition",     "") or "").strip()
    institution = (getattr(profile, "institution", "") or "").strip()
    isbn      = (getattr(profile, "isbn",        "") or "").strip()
    year      = getattr(profile, "copyright_year", None) or datetime.utcnow().year

    # vertical spacing
    for _ in range(4):
        doc.add_paragraph()

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run(title)
    r.bold = True
    r.font.name = _HEAD_FONT
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Subtitle
    if subtitle:
        ps = doc.add_paragraph()
        ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rs = ps.add_run(subtitle)
        rs.font.name = _HEAD_FONT
        rs.font.size = Pt(16)
        rs.italic = True

    doc.add_paragraph()

    # Author
    if author:
        pa = doc.add_paragraph()
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ra = pa.add_run(author)
        ra.font.name = _BODY_FONT
        ra.font.size = Pt(14)
        ra.bold = True

    if edition:
        pe = doc.add_paragraph()
        pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
        re_ = pe.add_run(edition)
        re_.font.name = _BODY_FONT
        re_.font.size = Pt(12)

    if institution:
        pi = doc.add_paragraph()
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ri = pi.add_run(institution)
        ri.font.name = _BODY_FONT
        ri.font.size = Pt(12)

    doc.add_paragraph()

    if isbn:
        pib = doc.add_paragraph()
        pib.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rib = pib.add_run(f"ISBN: {isbn}")
        rib.font.name = _BODY_FONT
        rib.font.size = Pt(10)
        rib.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    py_ = doc.add_paragraph()
    py_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ry = py_.add_run(str(year))
    ry.font.name = _BODY_FONT
    ry.font.size = Pt(11)
    ry.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    _add_hard_page_break(doc)


def _add_copyright_page(doc: DocxDocument, profile) -> None:
    author    = (getattr(profile, "author",      "") or "").strip()
    year      = getattr(profile, "copyright_year", None) or datetime.utcnow().year
    title     = (getattr(profile, "book_title",  "") or "").strip()
    custom    = (getattr(profile, "copyright_text", "") or "").strip()
    isbn      = (getattr(profile, "isbn",        "") or "").strip()
    institution = (getattr(profile, "institution", "") or "").strip()

    p_head = doc.add_paragraph()
    p_head.style = doc.styles["Front Matter Heading"]
    p_head.add_run("Copyright")

    if custom:
        for chunk in custom.split("\n\n"):
            cp = doc.add_paragraph(chunk.strip())
            cp.style = doc.styles["Body Text"]
    else:
        lines = [
            f"Copyright © {year} {author}." if author else f"Copyright © {year}.",
            "",
            "All rights reserved. No part of this publication may be reproduced, "
            "distributed, or transmitted in any form or by any means, including "
            "photocopying, recording, or other electronic or mechanical methods, "
            "without the prior written permission of the publisher.",
        ]
        if institution:
            lines.append(f"\nPublished by {institution}.")
        if isbn:
            lines.append(f"\nISBN: {isbn}")
        for line in lines:
            if line:
                doc.add_paragraph(line).style = doc.styles["Body Text"]
            else:
                doc.add_paragraph()

    _add_hard_page_break(doc)


def _add_front_matter_section(doc: DocxDocument, title: str, content: str) -> None:
    p_head = doc.add_paragraph()
    try:
        p_head.style = doc.styles["Front Matter Heading"]
    except Exception:
        p_head.style = doc.styles["Heading 1"]
    p_head.add_run(title)
    _page_break_before_para(p_head)

    if content and content.strip():
        for chunk in content.split("\n\n"):
            chunk = chunk.strip()
            if chunk:
                p = doc.add_paragraph(chunk)
                try:
                    p.style = doc.styles["Body Text"]
                except Exception:
                    p.style = doc.styles["Normal"]

    _add_hard_page_break(doc)


# ---------------------------------------------------------------------------
# Table styling
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color: str) -> None:
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs) -> None:
    """Set borders on a table cell. kwargs: top, bottom, left, right."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for side, color in kwargs.items():
        tag = OxmlElement(f"w:{side}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        tcBorders.append(tag)


def _add_table_to_doc(doc: DocxDocument, rows: List[List[str]]) -> None:
    """Add a professionally styled DOCX table from a list of row/cell strings."""
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    if num_cols == 0:
        return

    # Pad rows to uniform column count
    padded = [row + [""] * (num_cols - len(row)) for row in rows]

    tbl = doc.add_table(rows=len(padded), cols=num_cols)
    try:
        tbl.style = doc.styles["Table Grid"]
    except Exception:
        pass
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True

    border_color = "808080"

    for r_idx, row in enumerate(padded):
        is_header = (r_idx == 0)
        for c_idx, cell_text in enumerate(row):
            cell = tbl.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            if is_header:
                _shade_cell(cell, "2E2E2E")

            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(str(cell_text).strip())
            run.font.name = _BODY_FONT
            run.font.size = Pt(10)
            if is_header:
                run.bold = True
                run.font.color.rgb = _HEADER_FG

            _set_cell_border(
                cell,
                top=border_color, bottom=border_color,
                left=border_color, right=border_color,
            )

    # Add caption placeholder gap
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Content rendering
# ---------------------------------------------------------------------------

_H1_MARKERS = re.compile(r"^={3,}", re.MULTILINE)
_H2_MARKERS = re.compile(r"^-{3,}", re.MULTILINE)
_CHAPTER_PREFIX_RE = re.compile(r"^CHAPTER\s+\d+[:\-]?\s*", re.IGNORECASE)
_TABLE_SEP_RE = re.compile(r"\|")


def _parse_table_text(text: str) -> Optional[List[List[str]]]:
    """
    Parse a block of pipe-separated table text into row/cell lists.
    Returns None if the text does not look like a table.
    """
    lines = [l for l in text.splitlines() if "|" in l and l.strip()]
    if not lines:
        return None

    sep = " | " if " | " in lines[0] else "|"
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.split(sep)]
        cells = [c for c in cells if c]  # remove empty boundary cells
        if cells:
            rows.append(cells)
    return rows if rows else None


def _render_section_content(
    doc: DocxDocument,
    content: str,
    heading_mapping: Optional[dict],
    chapter_num: Optional[int] = None,
    chapter_title: Optional[str] = None,
    first_chapter: bool = False,
) -> None:
    """
    Render a single section's text content into the document.

    Supports:
    - H1/H2/H3 headings mapped via heading_mapping
    - Body paragraphs with justified alignment and first-line indent
    - Pipe-separated TABLE blocks → Word tables
    - Fenced heading markers (=== and ---) from the assembler
    - Chapter page-break logic
    """
    mapping = heading_mapping or {"H1": "Heading 1", "H2": "Heading 2", "H3": "Heading 3"}

    lines = content.splitlines()
    i = 0
    table_buffer: List[str] = []

    def _flush_table():
        nonlocal table_buffer
        if table_buffer:
            rows = _parse_table_text("\n".join(table_buffer))
            if rows:
                _add_table_to_doc(doc, rows)
            else:
                for tl in table_buffer:
                    _add_body_para(doc, tl)
            table_buffer = []

    def _add_body_para(doc: DocxDocument, text: str) -> None:
        p = doc.add_paragraph(text)
        try:
            p.style = doc.styles["Body Text"]
        except Exception:
            p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # If this section has a chapter number, add chapter opener heading
    if chapter_num is not None:
        _flush_table()
        opener_text = f"Chapter {chapter_num}"
        if chapter_title:
            title_clean = _CHAPTER_PREFIX_RE.sub("", chapter_title).strip()
            if title_clean:
                opener_text = f"Chapter {chapter_num}: {title_clean}"
        p_ch = doc.add_paragraph()
        try:
            p_ch.style = doc.styles["Chapter Opener"]
        except Exception:
            p_ch.style = doc.styles["Heading 1"]
        p_ch.add_run(opener_text)
        if not first_chapter:
            _page_break_before_para(p_ch)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip page markers
        if re.match(r"^=====\s*PAGE\s+\d+\s*=====", stripped):
            i += 1
            continue

        # Fenced H1 (=====...=====  or line followed by ====)
        if i + 1 < len(lines) and re.match(r"^={3,}$", lines[i + 1].strip()):
            _flush_table()
            if stripped:
                style_name = mapping.get("H1", "Heading 1")
                p = doc.add_paragraph(style=style_name)
                p.add_run(stripped)
                if chapter_num is None:
                    _page_break_before_para(p)
            i += 2
            continue

        # Fenced H2 (line followed by ---)
        if i + 1 < len(lines) and re.match(r"^-{3,}$", lines[i + 1].strip()):
            _flush_table()
            if stripped:
                style_name = mapping.get("H2", "Heading 2")
                p = doc.add_paragraph(style=style_name)
                p.add_run(stripped)
            i += 2
            continue

        # === separator lines themselves (already consumed above usually)
        if re.match(r"^={3,}$", stripped) or re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        # Table rows
        if "|" in stripped and stripped.count("|") >= 1:
            table_buffer.append(line)
            i += 1
            continue
        else:
            _flush_table()

        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        # Detect heading styles from the content itself (chapter heading prefix, all-caps short lines)
        if re.match(r"^CHAPTER\s+\d+", stripped, re.IGNORECASE):
            style_name = mapping.get("H1", "Heading 1")
            p = doc.add_paragraph(style=style_name)
            p.add_run(stripped)
            if chapter_num is None:
                _page_break_before_para(p)
            i += 1
            continue

        if re.match(r"^(PART|SECTION)\s+[IVXLCDM\d]+", stripped, re.IGNORECASE):
            style_name = mapping.get("H1", "Heading 1")
            p = doc.add_paragraph(style=style_name)
            p.add_run(stripped)
            if chapter_num is None:
                _page_break_before_para(p)
            i += 1
            continue

        # Sub-section numbered headings like "1.2 Heading"
        if re.match(r"^\d+\.\d+(\.\d+)?\s+\S", stripped):
            sub_depth = stripped.count(".")
            style_name = mapping.get(f"H{min(sub_depth + 1, 3)}", f"Heading {min(sub_depth + 1, 3)}")
            p = doc.add_paragraph(style=style_name)
            p.add_run(stripped)
            i += 1
            continue

        # Body paragraph
        _add_body_para(doc, stripped)
        i += 1

    _flush_table()


# ---------------------------------------------------------------------------
# Back matter
# ---------------------------------------------------------------------------

def _add_back_matter(doc: DocxDocument) -> None:
    """Add References and Index placeholder sections."""
    for section_title in ("References", "Index"):
        p = doc.add_paragraph()
        try:
            p.style = doc.styles["Heading 1"]
        except Exception:
            pass
        p.add_run(section_title)
        _page_break_before_para(p)

        placeholder = doc.add_paragraph(
            f"[{section_title} will be populated here.]"
        )
        try:
            placeholder.style = doc.styles["Body Text"]
        except Exception:
            pass
        _add_hard_page_break(doc)


# ---------------------------------------------------------------------------
# Main export function
# ---------------------------------------------------------------------------

def export_manuscript_docx(
    document,
    draft,
    profile,
    section_order: Optional[List[str]] = None,
) -> str:
    """
    Build a publication-quality DOCX from a ManuscriptDraft and ExportProfile.

    Returns the path to the saved .docx file.
    """
    doc = DocxDocument()

    # 1. Define all named styles
    _define_styles(doc)

    # 2. Page layout (A4, margins)
    _set_page_layout(doc, profile)

    # 3. Header + footer (before content so they apply to all pages)
    _add_header_footer(doc, profile)

    # 4. Title page (always)
    _add_title_page(doc, profile)

    # 5. Front matter sections
    fm_sections = getattr(profile, "front_matter_sections", None) or {}

    if fm_sections.get("copyright", True):
        _add_copyright_page(doc, profile)

    if fm_sections.get("dedication", False) and getattr(profile, "dedication", None):
        _add_front_matter_section(doc, "Dedication", profile.dedication)

    if fm_sections.get("preface", False) and getattr(profile, "preface", None):
        _add_front_matter_section(doc, "Preface", profile.preface)

    if fm_sections.get("acknowledgements", False) and getattr(profile, "acknowledgements", None):
        _add_front_matter_section(doc, "Acknowledgements", profile.acknowledgements)

    if fm_sections.get("disclaimer", False) and getattr(profile, "disclaimer", None):
        _add_front_matter_section(doc, "Disclaimer", profile.disclaimer)

    # 6. Table of Contents
    if getattr(profile, "include_toc", True):
        toc_heading = doc.add_paragraph()
        try:
            toc_heading.style = doc.styles["Front Matter Heading"]
        except Exception:
            toc_heading.style = doc.styles["Heading 1"]
        toc_heading.add_run("Table of Contents")
        _page_break_before_para(toc_heading)
        _add_toc(doc, profile)
        _add_hard_page_break(doc)

    # 7. Body sections
    sections = sorted(draft.sections, key=lambda s: s.section_order)
    if section_order:
        order_map = {sid: idx for idx, sid in enumerate(section_order)}
        sections = sorted(sections, key=lambda s: order_map.get(s.id, s.section_order))

    heading_mapping = getattr(profile, "heading_mapping", None) or {
        "H1": "Heading 1", "H2": "Heading 2", "H3": "Heading 3"
    }

    chapter_num = 0
    first_chapter_rendered = False

    for section in sections:
        current = section.current_version
        if not current:
            continue
        content = (current.content or "").strip()
        if not content:
            continue

        stype = section.section_type or "body"

        if stype == "front_matter":
            continue

        if stype in ("back_matter", "appendix"):
            # Appendix / back matter gets its own heading page
            p_app = doc.add_paragraph()
            try:
                p_app.style = doc.styles["Heading 1"]
            except Exception:
                pass
            p_app.add_run(section.title or stype.replace("_", " ").title())
            _page_break_before_para(p_app)
            _render_section_content(doc, content, heading_mapping)
            continue

        if stype == "chapter":
            chapter_num += 1
            _render_section_content(
                doc, content, heading_mapping,
                chapter_num=chapter_num,
                chapter_title=section.title,
                first_chapter=(not first_chapter_rendered),
            )
            first_chapter_rendered = True
            continue

        # Body section (not explicitly a chapter)
        _render_section_content(doc, content, heading_mapping)

    # 8. Back matter (References + Index placeholders)
    _add_back_matter(doc)

    # 9. Save
    storage_root = getattr(document, "storage_root", None) or os.path.join(
        os.path.dirname(document.local_storage_path), "exports"
    )
    export_dir = Path(storage_root) / "exports"
    export_dir.mkdir(parents=True, exist_ok=True)

    safe_title = (
        (getattr(profile, "book_title", "") or getattr(document, "filename", "manuscript") or "manuscript")
        .replace(" ", "_")[:50]
    )
    output_path = export_dir / f"{safe_title}_final.docx"
    doc.save(str(output_path))
    return str(output_path)
