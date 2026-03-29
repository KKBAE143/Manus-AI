#!/usr/bin/env python3

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

SOURCE_DIR = Path("final_clean_manuscript_v2")
OUTPUT_DIR = Path("ultra_final_manuscript")

DROP_PARAGRAPH_PATTERNS = [
    re.compile(r"let'?s make", re.I),
    re.compile(r"let'?s now", re.I),
    re.compile(r"we'?ll make", re.I),
    re.compile(r"we'?ll build", re.I),
    re.compile(r"we'?ll create", re.I),
    re.compile(r"we'?ll design", re.I),
    re.compile(r"we'?ll generate", re.I),
    re.compile(r"we'?ll now build", re.I),
    re.compile(r"we now begin", re.I),
    re.compile(r"we now move", re.I),
    re.compile(r"we now step", re.I),
    re.compile(r"we now continue", re.I),
    re.compile(r"here'?s how", re.I),
    re.compile(r"this section will", re.I),
    re.compile(r"this module will", re.I),
    re.compile(r"this part will", re.I),
    re.compile(r"this becomes", re.I),
    re.compile(r"this is where", re.I),
    re.compile(r"this is one of the most", re.I),
    re.compile(r"\bbuddy\b", re.I),
    re.compile(r"\bawesome\b", re.I),
    re.compile(r"\bperfect\b", re.I),
    re.compile(r"\bfantastic\b", re.I),
    re.compile(r"great job", re.I),
    re.compile(r"ready when you are", re.I),
    re.compile(r"say proceed", re.I),
    re.compile(r"once you confirm", re.I),
    re.compile(r"if you approve", re.I),
    re.compile(r"just say the word", re.I),
    re.compile(r"once you paste", re.I),
    re.compile(r"copy paste into word", re.I),
    re.compile(r"paste into word", re.I),
    re.compile(r"apply styles", re.I),
    re.compile(r"assign heading", re.I),
    re.compile(r"save as docx", re.I),
    re.compile(r"export to pdf", re.I),
    re.compile(r"in word assign heading", re.I),
    re.compile(r"apply heading style", re.I),
    re.compile(r"paste this section", re.I),
    re.compile(r"copy this into word", re.I),
    re.compile(r"word file", re.I),
    re.compile(r"layout selection", re.I),
    re.compile(r"title page discussion", re.I),
    re.compile(r"cover page discussion", re.I),
    re.compile(r"branding section", re.I),
    re.compile(r"manual structure planning", re.I),
    re.compile(r"format disorder", re.I),
    re.compile(r"format section", re.I),
    re.compile(r"format module", re.I),
    re.compile(r"proceed to disorder", re.I),
    re.compile(r"disorder completed", re.I),
    re.compile(r"completed successfully", re.I),
    re.compile(r"ready for next", re.I),
    re.compile(r"next module", re.I),
    re.compile(r"next section preview", re.I),
    re.compile(r"\[PAGE BREAK\]", re.I),
    re.compile(r"\[PASTE START\]", re.I),
    re.compile(r"\[PASTE END\]", re.I),
]

DROP_SENTENCE_PATTERNS = DROP_PARAGRAPH_PATTERNS

STAR_RE = re.compile(r"[⭐]+")
MULTIBLANK_RE = re.compile(r"\n{3,}")


def iter_blocks(doc: DocxDocumentType):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def normalize(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = STAR_RE.sub("", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def should_drop_paragraph(text: str) -> bool:
    candidate = normalize(text)
    if not candidate:
        return False
    return any(pattern.search(candidate) for pattern in DROP_PARAGRAPH_PATTERNS)


def clean_text(text: str) -> str:
    text = text.replace("format Disorder 48 - Recurrent Urinary Tract Infection", "DISORDER 48 - Recurrent Urinary Tract Infection")
    text = STAR_RE.sub("", text)
    lines = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            lines.append("")
            continue
        if line.lower().startswith("[source page "):
            continue
        if any(pattern.search(line) for pattern in DROP_PARAGRAPH_PATTERNS):
            continue
        lines.append(line)

    text = "\n".join(lines)
    text = MULTIBLANK_RE.sub("\n\n", text).strip()
    pieces = re.split(r"(?<=[.!?])\s+", text)
    kept = []
    for piece in pieces:
        candidate = normalize(piece)
        if not candidate:
            continue
        if any(pattern.search(candidate) for pattern in DROP_SENTENCE_PATTERNS):
            continue
        kept.append(candidate)
    return "\n".join(kept).strip()


def copy_table(dst: DocxDocumentType, table: Table) -> None:
    cleaned_rows = []
    for row in table.rows:
        values = [clean_text(cell.text) for cell in row.cells]
        if any(value.strip() for value in values):
            cleaned_rows.append(values)
    if not cleaned_rows:
        return
    new_table = dst.add_table(rows=0, cols=len(cleaned_rows[0]))
    for values in cleaned_rows:
        row = new_table.add_row().cells
        for idx, value in enumerate(values):
            row[idx].text = value


def clean_file(input_path: Path, output_path: Path) -> None:
    src = Document(str(input_path))
    dst = Document()
    for block in iter_blocks(src):
        if isinstance(block, Paragraph):
            if should_drop_paragraph(block.text):
                continue
            cleaned = clean_text(block.text)
            if cleaned:
                dst.add_paragraph(cleaned)
        else:
            copy_table(dst, block)
    dst.save(str(output_path))


def main() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    for part in range(1, 16):
        name = f"textbook_part_{part:02d}.docx"
        clean_file(SOURCE_DIR / name, OUTPUT_DIR / name)
    print("Ultra final manuscript complete.")


if __name__ == "__main__":
    main()
